import tempfile
import unittest
from pathlib import Path
from unittest.mock import PropertyMock, patch

from docx import Document
from PIL import Image, ImageDraw
from openpyxl import Workbook

from app.agents.document_parser import (
    DocumentParserAgent,
    _clean_field_value,
    _expected_visual_marks,
    _extract_fields,
    _infer_subtype_from_content,
    _response_entities,
    _response_tenderer_field,
    _seal_matches_entity,
)
from app.api.file_helpers import infer_file_type
from app.schemas.task import UploadedFileInfo
from app.services.file_parser import ParsedFileContent, ParsedPage, ParsedTableData, _merge_continued_tables, document_tool_registry, parse_file
from app.services.document_semantic_enhancer import (
    DifyWorkflowError,
    DocumentSemanticEnhancer,
    document_semantic_enhancer,
)


SAMPLE_TEXT = """项目名称：某市信息化平台建设项目
项目预算：人民币100万元
最高投标限价：人民币80万元
采购人：某市信息中心

一、投标人资格要求
投标人应当具有独立承担民事责任的能力。

二、评审办法
技术方案评分40分。
"""


class FieldCleanupTests(unittest.TestCase):
    def test_repeated_label_and_inline_next_field_are_removed(self) -> None:
        self.assertEqual(
            _clean_field_value("project_name", "项目名称XX市信息化平台升级建设项目"),
            "XX市信息化平台升级建设项目",
        )
        self.assertEqual(
            _clean_field_value(
                "tenderer",
                "XX市政务服务管理中心。采购代理机构：XX市公共资源交易中心",
            ),
            "XX市政务服务管理中心",
        )


class FileParserTests(unittest.TestCase):
    def test_parse_utf8_text(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "sample.txt"
            path.write_text(SAMPLE_TEXT, encoding="utf-8")

            parsed = parse_file(path)

            self.assertIn("某市信息化平台", parsed.text)
            self.assertEqual(parsed.page_count, 1)
            self.assertFalse(parsed.is_scanned)
            self.assertEqual(parsed.selected_tool, "通用文本解析工具")
            self.assertTrue(parsed.tool_trace)

    def test_parse_docx_with_table(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "sample.docx"
            document = Document()
            document.add_paragraph("项目名称：文档解析测试项目")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "字段"
            table.cell(0, 1).text = "值"
            table.cell(1, 0).text = "预算"
            table.cell(1, 1).text = "100万元"
            document.save(path)

            parsed = parse_file(path)

            self.assertEqual(len(parsed.tables), 1)
            self.assertIn("预算 | 100万元", parsed.text)
            self.assertEqual(parsed.selected_tool, "DOCX段落与表格解析工具")
            self.assertTrue(any(item.element_type == "table" for item in parsed.layout_elements))

    def test_merges_same_header_tables_on_consecutive_pages(self) -> None:
        tables = [
            ParsedTableData(page=1, page_end=1, rows=[["投标人", "得分"], ["甲公司", "88"]]),
            ParsedTableData(page=2, page_end=2, rows=[["投标人", "得分"], ["乙公司", "86"]]),
            ParsedTableData(page=4, page_end=4, rows=[["投标人", "报价"], ["丙公司", "90万元"]]),
        ]
        merged = _merge_continued_tables(tables)
        self.assertEqual(len(merged), 2)
        self.assertTrue(merged[0].continued)
        self.assertEqual(merged[0].page_end, 2)
        self.assertEqual(merged[0].rows[-1], ["乙公司", "86"])

    def test_tool_registry_capabilities(self) -> None:
        capabilities = document_tool_registry.capabilities()
        names = {item["name"] for item in capabilities}
        self.assertEqual(
            names,
            {
                "通用文本解析工具",
                "PDF文本与表格解析工具",
                "DOCX段落与表格解析工具",
                "RapidOCR扫描PDF识别工具",
                "XLSX评审数据解析工具",
            },
        )

    def test_parse_realistic_evaluation_xlsx(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "专家评分汇总表.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "评分汇总表"
            sheet.append(["投标人", "评审专家", "评审因素", "满分", "得分", "权重", "折算得分", "总得分", "排名", "投标报价"])
            sheet.append(["甲科技有限公司", "专家A", "技术方案", 40, 36, 0.4, 14.4, 88.5, 1, 980000])
            sheet.append(["乙信息有限公司", "专家A", "技术方案", 40, 34, 0.4, 13.6, 85.2, 2, 990000])
            workbook.save(path)
            workbook.close()

            parsed = parse_file(path)
            self.assertEqual(parsed.sheet_names, ["评分汇总表"])
            self.assertEqual(parsed.tables[0].sheet, "评分汇总表")
            file_info = UploadedFileInfo(file_id="X001", filename=path.name, file_type="评审评分表", saved_path=str(path))
            documents, _, _ = DocumentParserAgent().run([file_info], "测试项目")
            document = documents[0]
            self.assertEqual(document.document_subtype, "评分汇总表")
            self.assertEqual(len(document.score_details), 2)
            self.assertEqual(document.score_details[0].factor, "技术方案")
            self.assertEqual(document.score_details[0].source.row, 2)
            self.assertEqual(document.opening_records[0].bid_price, 980000)
            self.assertEqual(document.score_summaries[0].rank, 1)
            self.assertEqual(document.candidate_rankings[0].bidder, "甲科技有限公司")

    def test_scanned_pdf_uses_ocr_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "scanned.pdf"
            image = Image.new("RGB", (1000, 180), "white")
            ImageDraw.Draw(image).text(
                (30, 50),
                "Tender OCR Project 123",
                fill="black",
                font_size=48,
            )
            image.save(path, "PDF", resolution=150)
            image.close()

            parsed = parse_file(path)

            self.assertTrue(parsed.is_scanned)
            self.assertTrue(parsed.ocr_applied)
            self.assertGreater(parsed.ocr_confidence, 0.7)
            self.assertIn("RapidOCR", parsed.selected_tool)
            self.assertIn("123", parsed.text)


class DocumentParserAgentTests(unittest.TestCase):
    def test_cover_title_is_used_but_institution_name_is_not_project_name(self) -> None:
        content = ParsedFileContent(
            text="武汉大学\n某市政务信息化平台运行维护服务项目\n投标文件\n",
            pages=[ParsedPage(number=1, text="武汉大学\n某市政务信息化平台运行维护服务项目\n投标文件\n")],
        )
        fields = _extract_fields(content, "武汉大学")
        self.assertEqual(fields["project_name"].value, "某市政务信息化平台运行维护服务项目")
        self.assertEqual(fields["project_name"].source_location, "第 1 页封面标题")
        self.assertGreaterEqual(fields["project_name"].confidence, 0.75)

    def test_generic_real_response_document_recovers_entities(self) -> None:
        text = """响 应 文 件
科研委外合作项目研究响应文件
合 作 方 ： 武汉大学
致中国烟草总公司湖北省公司信息中心:
承诺单位：武汉大学（盖章）
不得与采购人、其他供应商或者采购代理机构恶意串通。
"""
        self.assertEqual(_infer_subtype_from_content("业务文件", text), "响应文件")
        self.assertEqual(_response_entities(text), ["武汉大学"])
        field = _response_tenderer_field(
            ParsedFileContent(text=text, pages=[ParsedPage(number=1, text=text)])
        )
        self.assertIsNotNone(field)
        self.assertEqual(field.value, "中国烟草总公司湖北省公司信息中心")
        self.assertTrue(field.requires_human_review)

    def test_identifies_expected_seal_and_signature_positions(self) -> None:
        content = ParsedFileContent(
            text="投标人（盖章）：\n法定代表人（签字）：",
            pages=[ParsedPage(number=3, text="投标人（盖章）：\n法定代表人（签字）：")],
        )
        requirements = _expected_visual_marks(content)
        self.assertIn(("seal", 3, "投标人（盖章）："), requirements)
        self.assertIn(("signature", 3, "法定代表人（签字）："), requirements)

    def test_matches_recognized_seal_text_to_document_entity(self) -> None:
        self.assertTrue(_seal_matches_entity("甲科技有限公司", ["投标人：甲科技有限公司"]))
        self.assertTrue(_seal_matches_entity("诚科技有限公司\n系统测试专用章", ["华诚科技有限公司"]))
        self.assertFalse(_seal_matches_entity("乙信息有限公司", ["投标人：甲科技有限公司"]))
        self.assertFalse(_seal_matches_entity("成科技有限公司\n系统测试专用章", ["天远科技有限公司"]))
        self.assertIsNone(_seal_matches_entity("", ["投标人：甲科技有限公司"]))

    def test_extracts_rejection_opinion_and_candidate_ranking(self) -> None:
        text = """项目名称：某信息化运维项目
评审结论：经评审委员会评议，推荐以下中标候选人。
第一中标候选人：甲科技有限公司。
投标人：乙信息有限公司，否决投标理由：未按要求提交投标保证金。
"""
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "评标报告.txt"
            path.write_text(text, encoding="utf-8")
            info = UploadedFileInfo(file_id="R001", filename=path.name, file_type="评标报告", saved_path=str(path))
            documents, _, _ = DocumentParserAgent().run([info], "测试项目")
            document = documents[0]
            self.assertEqual(document.document_subtype, "评标报告")
            self.assertEqual(document.evaluation_opinions[0].opinion, "经评审委员会评议，推荐以下中标候选人。")
            self.assertEqual(document.candidate_rankings[0].bidder, "甲科技有限公司")
            self.assertEqual(document.rejection_records[0].bidder, "乙信息有限公司")
            self.assertIn("未按要求提交投标保证金", document.rejection_records[0].reason)

    def test_procurement_invalid_response_rules_are_not_rejection_records(self) -> None:
        text = """项目名称：某采购项目
第一章 响应规则
供应商未按要求提交保证金的，响应无效。
投标报价超过最高投标限价的，投标无效。
"""
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "采购文件.txt"
            path.write_text(text, encoding="utf-8")
            info = UploadedFileInfo(file_id="R002", filename=path.name, file_type="招标文件", saved_path=str(path))
            documents, _, _ = DocumentParserAgent().run([info], "测试项目")

        document = documents[0]
        self.assertEqual(document.rejection_records, [])
        self.assertEqual(len(document.invalid_bid_clauses), 2)
        self.assertTrue(any("响应无效" in item for item in document.invalid_bid_clauses))

    def test_structured_document_and_quality_checks(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "招标文件.txt"
            path.write_text(SAMPLE_TEXT, encoding="utf-8")
            file_info = UploadedFileInfo(
                file_id="F001",
                filename=path.name,
                file_type="招标文件",
                saved_path=str(path),
            )

            documents, result, raw_texts = DocumentParserAgent().run(
                [file_info],
                "备用项目名称",
            )

            self.assertEqual(len(documents), 1)
            document = documents[0]
            self.assertEqual(document.project_name, "某市信息化平台建设项目")
            self.assertIn("budget", document.extracted_fields)
            self.assertIn("price_limit", document.extracted_fields)
            self.assertEqual(document.selected_tool, "通用文本解析工具")
            self.assertTrue(any(section.title == "一、投标人资格要求" for section in document.sections))
            self.assertTrue(all(check.status == "passed" for check in document.quality_checks))
            self.assertFalse(result.data["requires_human_review"])
            self.assertIn("F001", raw_texts)

    def test_numbered_clauses_are_content_and_procurement_unit_is_tenderer(self) -> None:
        text = """第一章 项目基本情况
项目名称：某市智慧政务服务平台升级项目。
采购单位：某市政务服务管理中心。

第二章 投标人资格要求
1. 投标人应当具有独立承担民事责任的能力。
2. 投标人注册地址必须位于本市。

第三章 评审办法
1. 评审内容包括投标报价和技术方案。
"""
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "测试招标文件.txt"
            path.write_text(text, encoding="utf-8")
            file_info = UploadedFileInfo(
                file_id="F002",
                filename=path.name,
                file_type="招标文件",
                saved_path=str(path),
            )

            with patch.object(
                DocumentSemanticEnhancer,
                "enabled",
                new_callable=PropertyMock,
                return_value=False,
            ):
                documents, _, _ = DocumentParserAgent().run([file_info], "备用项目")

            document = documents[0]
            self.assertEqual(document.tenderer, "某市政务服务管理中心")
            titles = [section.title for section in document.sections]
            self.assertEqual(
                titles,
                ["第一章 项目基本情况", "第二章 投标人资格要求", "第三章 评审办法"],
            )
            qualification = next(
                section for section in document.sections if section.title == "第二章 投标人资格要求"
            )
            self.assertIn("1. 投标人应当具有独立承担民事责任的能力。", qualification.content)
            self.assertIn("2. 投标人注册地址必须位于本市。", qualification.content)

    def test_procurement_agency_is_not_used_as_tenderer(self) -> None:
        text = """项目名称：某建设项目
招标人：某大学
招标代理机构：某工程咨询有限公司
"""
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "评标报告.txt"
            path.write_text(text, encoding="utf-8")
            info = UploadedFileInfo(file_id="AGENCY-1", filename=path.name, file_type="评标报告", saved_path=str(path))
            with patch.object(DocumentSemanticEnhancer, "enabled", new_callable=PropertyMock, return_value=False):
                documents, _, _ = DocumentParserAgent().run([info], "备用项目")

        self.assertEqual(documents[0].tenderer, "某大学")
        self.assertEqual(documents[0].procurement_agency, "某工程咨询有限公司")

    def test_template_placeholders_do_not_consume_following_lines(self) -> None:
        text = """北京市政府采购项目
公开招标文件示范文本
项目名称：
项目编号/包号：
采 购 人：
采购代理机构：
使用说明
请各预算单位、采购代理机构认真组织示范文本的推广使用。
项目预算金额：____万元、项目最高限价（如有）：____万元
采购代理机构：指依法进行政府采购的代理组织，本项目采购人见第一章。
最高限价（如有）：____万元
提交投标文件的截止时间和开标时间。
第一章 投标邀请
"""
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "示范文本.txt"
            path.write_text(text, encoding="utf-8")
            file_info = UploadedFileInfo(
                file_id="F003",
                filename=path.name,
                file_type="招标文件",
                saved_path=str(path),
            )

            with patch.object(
                DocumentSemanticEnhancer,
                "enabled",
                new_callable=PropertyMock,
                return_value=False,
            ):
                documents, _, _ = DocumentParserAgent().run([file_info], "示范文本解析测试")

            document = documents[0]
            self.assertEqual(document.project_name, "示范文本解析测试")
            self.assertEqual(document.tenderer, "")
            self.assertNotIn("budget", document.extracted_fields)
            self.assertNotIn("price_limit", document.extracted_fields)
            self.assertNotIn("deadline", document.extracted_fields)
            self.assertTrue(document.extracted_fields["project_name"].requires_human_review)

    def test_evaluation_report_deadline_sentence_is_extracted(self) -> None:
        text = "截止至本项目投标文件截止时间 2025 年 3 月 19 日 9 时 00 分共有7家单位递交投标文件。"
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "评标报告.txt"
            path.write_text(text, encoding="utf-8")
            file_info = UploadedFileInfo(
                file_id="F-DEADLINE",
                filename=path.name,
                file_type="评标报告",
                saved_path=str(path),
            )
            with patch.object(
                DocumentSemanticEnhancer,
                "enabled",
                new_callable=PropertyMock,
                return_value=False,
            ):
                documents, _, _ = DocumentParserAgent().run([file_info], "截止时间测试")

        deadline = documents[0].extracted_fields["deadline"]
        self.assertIn("2025 年 3 月 19 日 9 时 00 分", deadline.value)

    def test_unfilled_template_skips_dify_semantic_enhancement(self) -> None:
        text = """公开招标文件示范文本
使用说明
二、填写规则
项目名称：____
项目预算：____万元
最高投标限价：____万元
采购人：____
第一章 投标邀请
"""
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "公开招标文件示范文本.txt"
            path.write_text(text, encoding="utf-8")
            file_info = UploadedFileInfo(
                file_id="F-TEMPLATE",
                filename=path.name,
                file_type="招标文件",
                saved_path=str(path),
            )
            with patch.object(
                DocumentSemanticEnhancer,
                "enabled",
                new_callable=PropertyMock,
                return_value=True,
            ), patch(
                "app.agents.document_parser.document_semantic_enhancer.enhance"
            ) as semantic_enhance:
                documents, result, _ = DocumentParserAgent().run([file_info], "示范模板测试")

            semantic_enhance.assert_not_called()
            document = documents[0]
            self.assertTrue(any("跳过 Dify" in item for item in document.warnings))
            self.assertTrue(result.data["requires_human_review"])
            self.assertTrue(result.issues)
            self.assertTrue(
                all(issue.assessment == "待人工判断" for issue in result.issues)
            )
            self.assertTrue(
                any(issue.issue_type == "关键字段待确认" for issue in result.issues)
            )
            self.assertTrue(
                any("低置信度字段" in issue.description for issue in result.issues)
            )

    def test_wrapped_pdf_project_name_is_joined(self) -> None:
        text = """项目名称：北京市流动人员人事档案公共服务管理子系统运行保
障服务
项目编号/包号：0733-25183857/01
第一章 采购邀请
"""
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "采购文件.txt"
            path.write_text(text, encoding="utf-8")
            file_info = UploadedFileInfo(
                file_id="F004",
                filename=path.name,
                file_type="招标文件",
                saved_path=str(path),
            )
            with patch.object(
                DocumentSemanticEnhancer,
                "enabled",
                new_callable=PropertyMock,
                return_value=False,
            ):
                documents, _, _ = DocumentParserAgent().run([file_info], "备用项目")

            self.assertEqual(
                documents[0].project_name,
                "北京市流动人员人事档案公共服务管理子系统运行保障服务",
            )

    def test_procurement_document_filename_is_tender_document(self) -> None:
        self.assertEqual(infer_file_type("某项目采购文件.pdf"), "招标文件")
        self.assertEqual(infer_file_type("某项目竞争性磋商文件.docx"), "招标文件")


class DocumentSemanticEnhancerTests(unittest.TestCase):
    def test_candidate_ranking_requires_positive_explicit_rank(self) -> None:
        payload = {
            "sections": [],
            "key_fields": [],
            "candidate_rankings": [
                {"bidder": "甲公司", "rank": 0, "evidence": "甲公司 1000000"},
                {"bidder": "乙公司", "rank": 1, "evidence": "合格的中标候选人名称：乙公司"},
            ],
        }
        enhancer = DocumentSemanticEnhancer()
        source = "甲公司 1000000\n合格的中标候选人名称：乙公司"
        with patch(
            "app.services.document_semantic_enhancer.dify_client.run_document_semantic_parser",
            return_value=payload,
        ):
            result = enhancer.enhance(
                source,
                "F-NO-RANK",
                requested_fields=["candidate_rankings"],
                include_sections=False,
                document_type="评标报告",
            )

        self.assertEqual(result.candidate_rankings, [])

    def test_uses_new_five_input_contract_and_parses_incremental_records(self) -> None:
        payload = {
            "document_type": "评标报告",
            "sections": [],
            "key_fields": [{
                "field_name": "project_name",
                "value": "测试项目",
                "raw_text": "项目名称：测试项目。",
                "confidence": 0.95,
                "requires_human_review": False,
            }],
            "rejection_records": [],
            "evaluation_opinions": [],
            "candidate_rankings": [{
                "bidder": "甲公司",
                "rank": 1,
                "lot": "",
                "evidence": "第一中标候选人：甲公司。",
            }],
            "warnings": [],
        }
        text = "项目名称：测试项目。\n第一中标候选人：甲公司。"
        enhancer = DocumentSemanticEnhancer()
        with patch(
            "app.services.document_semantic_enhancer.dify_client.run_document_semantic_parser",
            return_value=payload,
        ) as mocked:
            result = enhancer.enhance(
                text,
                "F-V2",
                requested_fields=["project_name", "candidate_rankings"],
                include_sections=False,
                document_type="评标报告",
                parser_context={"key_fields": {}},
            )

        args = mocked.call_args.args
        self.assertEqual(args[1], "评标报告")
        self.assertIn('"key_fields"', args[2])
        self.assertIn('"candidate_rankings"', args[3])
        self.assertEqual(args[4], "false")
        self.assertEqual(result.fields["project_name"].value, "测试项目")
        self.assertEqual(result.candidate_rankings[0].bidder, "甲公司")

    def test_parse_dify_semantic_result(self) -> None:
        payload = {
            "sections": [
                {
                    "title": "投标人资格要求",
                    "level": 1,
                    "content": "投标人应具备相应能力。",
                    "page": 3,
                }
            ],
            "key_fields": {
                "budget": {
                    "value": "100万元",
                    "raw_text": "项目预算：100万元",
                    "source_location": "第2页",
                    "confidence": 0.88,
                    "requires_human_review": False,
                },
                "price_limit": {
                    "value": "None",
                    "raw_text": "最高限价：____万元",
                    "confidence": 0.1,
                },
            },
            "warnings": ["字段来自语义增强"],
        }
        enhancer = DocumentSemanticEnhancer()
        with patch(
            "app.services.document_semantic_enhancer.dify_client.run_document_semantic_parser",
            return_value=payload,
        ):
            sections, fields, warnings = enhancer.enhance(
                "项目预算：100万元\n投标人应具备相应能力。",
                "F001",
            )

        self.assertEqual(sections[0].title, "投标人资格要求")
        self.assertEqual(fields["budget"].value, "100万元")
        self.assertNotIn("price_limit", fields)
        self.assertAlmostEqual(fields["budget"].confidence, 0.88)
        self.assertEqual(warnings, ["字段来自语义增强"])

    def test_semantic_price_limit_must_be_supported_by_limit_evidence(self) -> None:
        payload = {
            "sections": [],
            "key_fields": {
                "price_limit": {
                    "value": "249万元",
                    "raw_text": "项目预算金额：249万元",
                    "confidence": 0.95,
                }
            },
        }
        enhancer = DocumentSemanticEnhancer()
        with patch(
            "app.services.document_semantic_enhancer.dify_client.run_document_semantic_parser",
            return_value=payload,
        ):
            _, fields, _ = enhancer.enhance(
                "项目预算金额：249万元",
                "F-PRICE",
                requested_fields=["price_limit"],
                include_sections=False,
            )

        self.assertNotIn("price_limit", fields)

    def test_long_document_is_chunked_and_fields_are_merged(self) -> None:
        enhancer = DocumentSemanticEnhancer()
        long_text = ("A" * 5900) + "\n" + ("B" * 1000)
        responses = [
            {
                "sections": [{"title": "第一章", "content": "内容一"}],
                "key_fields": {
                    "project_name": {"value": "测试项目", "confidence": 0.8}
                },
            },
            {
                "sections": [{"title": "第二章", "content": "内容二"}],
                "key_fields": {
                    "budget": {"value": "100万元", "confidence": 0.9}
                },
            },
        ]
        with patch(
            "app.services.document_semantic_enhancer.dify_client.run_document_semantic_parser",
            side_effect=responses,
        ) as mocked:
            sections, fields, warnings = enhancer.enhance(long_text, "F-LONG")

        self.assertEqual(mocked.call_count, 2)
        self.assertEqual([item.title for item in sections], ["第一章", "第二章"])
        self.assertEqual(fields["project_name"].value, "测试项目")
        self.assertEqual(fields["budget"].value, "100万元")
        self.assertEqual(warnings, [])

    def test_failed_chunk_is_bisected_and_retried(self) -> None:
        enhancer = DocumentSemanticEnhancer()
        text = ("第一部分内容。\n" * 300).strip()
        responses = [
            DifyWorkflowError("输出 JSON 被截断"),
            {"sections": [], "key_fields": {"budget": {"value": "100万元"}}},
            {"sections": [], "key_fields": {"tenderer": {"value": "某采购人"}}},
        ]
        with patch(
            "app.services.document_semantic_enhancer.dify_client.run_document_semantic_parser",
            side_effect=responses,
        ) as mocked:
            _, fields, warnings = enhancer.enhance(text, "F-RETRY")

        self.assertEqual(mocked.call_count, 3)
        self.assertEqual(fields["budget"].value, "100万元")
        self.assertEqual(fields["tenderer"].value, "某采购人")
        self.assertTrue(any("缩小分段重试" in item for item in warnings))

    def test_transport_failure_does_not_bisect_and_retry(self) -> None:
        enhancer = DocumentSemanticEnhancer()
        text = ("评标报告内容。\n" * 400).strip()
        with patch(
            "app.services.document_semantic_enhancer.dify_client.run_document_semantic_parser",
            side_effect=DifyWorkflowError("Server Unavailable: failed to resolve api.example.com"),
        ) as mocked:
            with self.assertRaises(DifyWorkflowError):
                enhancer.enhance(text, "F-NETWORK")

        self.assertEqual(mocked.call_count, 1)

    def test_conflicting_field_values_require_human_review(self) -> None:
        enhancer = DocumentSemanticEnhancer()
        long_text = ("A" * 5900) + "\n" + ("B" * 1000)
        responses = [
            {"sections": [], "key_fields": {"budget": {"value": "100万元", "confidence": 0.8}}},
            {"sections": [], "key_fields": {"budget": {"value": "120万元", "confidence": 0.9}}},
        ]
        with patch(
            "app.services.document_semantic_enhancer.dify_client.run_document_semantic_parser",
            side_effect=responses,
        ):
            _, fields, warnings = enhancer.enhance(long_text, "F-CONFLICT")

        self.assertEqual(fields["budget"].value, "120万元")
        self.assertTrue(fields["budget"].requires_human_review)
        self.assertTrue(any("识别结果不一致" in item for item in warnings))

    def test_missing_field_only_selects_relevant_chunks(self) -> None:
        enhancer = DocumentSemanticEnhancer()
        long_text = "\n".join(
            [
                "项目背景说明" * 700,
                "技术服务要求" * 700,
                "最高投标限价：200万元。" + ("报价说明" * 600),
                "售后服务要求" * 700,
                "合同履约要求" * 700,
            ]
        )
        response = {
            "sections": [],
            "key_fields": {
                "price_limit": {
                    "value": "200万元",
                    "raw_text": "最高投标限价：200万元。",
                    "confidence": 0.95,
                },
                "project_name": {"value": "不应合并的项目名称", "confidence": 0.99},
            },
            "warnings": ["未找到最高投标限价字段", "未找到项目名称字段"],
        }
        with patch(
            "app.services.document_semantic_enhancer.dify_client.run_document_semantic_parser",
            return_value=response,
        ) as mocked:
            _, fields, warnings = enhancer.enhance(
                long_text,
                "F-TARGETED",
                requested_fields=["price_limit"],
                include_sections=False,
            )

        total_chunks = len(enhancer._split_text(long_text))
        self.assertLess(mocked.call_count, total_chunks)
        self.assertEqual(fields["price_limit"].value, "200万元")
        self.assertNotIn("project_name", fields)
        self.assertEqual(warnings, ["未找到最高投标限价字段"])


if __name__ == "__main__":
    unittest.main()
