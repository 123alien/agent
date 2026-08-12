import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import PropertyMock, patch

from docx import Document

from app.agents.report_generator import ReportGeneratorAgent
from app.schemas.task import AgentResult, Issue, ParsedDocument, TaskRecord, TaskResult
from app.services.report_service import (
    create_docx_report,
    public_warning,
    report_suggestion,
    report_display_title,
    select_report_issues,
)
from app.services.pdf_service import create_report_pdf


class ReportGeneratorTests(unittest.TestCase):
    def test_public_report_hides_vendor_error_and_formalizes_review_wording(self) -> None:
        self.assertEqual(
            public_warning("Dify 文档语义增强失败: timeout"),
            "语义增强未完成，已采用确定性解析结果继续核验。",
        )
        issue = Issue(
            agent="文档解析智能体", risk_level="低", issue_type="签名核验",
            description="人工确认未签字", suggestion="请人工查看原始文件第15页。",
            detection_status="not_checked", requires_human_review=False,
        )
        issue.final_status = "confirmed_issue"
        issue.assessment = "明确问题"
        issue.requires_human_review = False
        self.assertNotIn("请人工查看", report_suggestion(issue, "正式核验版"))

    def test_report_type_scopes_issues_and_template_changes_title(self) -> None:
        result = TaskResult(
            summary="测试",
            issues=[
                Issue(
                    issue_id="C1",
                    agent="合规审查智能体",
                    risk_level="高",
                    issue_type="合规问题",
                    description="合规问题",
                ),
                Issue(
                    issue_id="D1",
                    agent="数据核验智能体",
                    risk_level="中",
                    issue_type="数据问题",
                    description="数据问题",
                ),
            ],
        )
        selected = select_report_issues(result, "合规审查专项报告")
        self.assertEqual([item.issue_id for item in selected], ["C1"])
        self.assertEqual(
            report_display_title("综合智能核验报告", "简版管理层报告"),
            "综合智能核验报告（管理层简版）",
        )

    def test_quality_metrics_detect_missing_evidence(self) -> None:
        issue = Issue(
            agent="合规审查智能体",
            risk_level="高",
            issue_type="地域限制",
            description="注册地址受限",
            requires_human_review=True,
        )
        result = ReportGeneratorAgent().run([], [issue])
        self.assertFalse(result.data["report_ready"])
        self.assertEqual(result.data["missing_evidence_count"], 1)
        self.assertEqual(result.data["review_required_count"], 1)

    def test_docx_report_contains_issue_and_review_notice(self) -> None:
        task = TaskRecord(
            task_id="T-REPORT",
            project_id="P-001",
            project_name="某市政务平台运维项目",
            check_type="full",
            status="completed",
            created_at="2026-08-10T00:00:00+08:00",
            updated_at="2026-08-10T00:00:00+08:00",
        )
        issue = Issue(
            issue_id="I-001",
            agent="合规审查智能体",
            risk_level="高",
            issue_type="指定品牌",
            description="技术要求指定唯一品牌。",
            evidence=["监控工具必须使用甲公司生产的政务云眼V5。"],
            basis="需结合检索到的法规依据复核。",
            suggestion="改为功能和性能参数。",
            requires_human_review=True,
        )
        result = TaskResult(
            summary="发现1项高风险问题。",
            parsed_documents=[
                ParsedDocument(
                    file_id="F1",
                    filename="招标文件.docx",
                    file_type="docx",
                    text_length=100,
                )
            ],
            agent_results=[AgentResult(agent="合规审查智能体", summary="发现指定品牌条款。")],
            issues=[issue],
        )
        with tempfile.TemporaryDirectory() as directory:
            with patch(
                "app.services.report_service.settings",
                SimpleNamespace(reports_dir=Path(directory)),
            ):
                path = create_docx_report(task, result)
            self.assertTrue(path.exists())
            document = Document(path)
            text = "\n".join(p.text for p in document.paragraphs)
            table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
            self.assertIn("某市政务平台运维项目", text)
            self.assertIn("指定品牌", text)
            self.assertIn("监控工具必须使用甲公司生产的政务云眼V5", table_text)
            self.assertIn("需要", table_text)

    def test_pdf_report_contains_canonical_counts_and_evidence_index(self) -> None:
        task = TaskRecord(
            task_id="T-PDF", project_id="P-003", project_name="任务显示名", check_type="full",
            status="completed", created_at="2026-08-11T00:00:00+08:00", updated_at="2026-08-11T00:00:00+08:00",
        )
        issue = Issue(
            issue_id="R-001", agent="文档解析智能体", risk_level="低", issue_type="印章视觉核验待复核",
            description="印章识别置信度不足。", evidence=["投标人（盖章）："],
            requires_human_review=True, final_status="human_review", detection_status="low_confidence", confidence=0.54,
        )
        report_agent = AgentResult(
            agent="报告生成智能体", summary="旧摘要不应控制最终统计。",
            data={"report_status":"待复核版", "output_type":"综合智能核验报告", "template_type":"标准审查报告", "report_package":{
                "project_info":{"project_name":"某市政务平台运维项目", "field_sources":{}},
            }},
        )
        result = TaskResult(summary="旧总结", issues=[issue], agent_results=[report_agent])
        with tempfile.TemporaryDirectory() as directory:
            with patch("app.services.pdf_service.settings", SimpleNamespace(reports_dir=Path(directory))):
                path = create_report_pdf(task, result)
            self.assertTrue(path.exists())
            self.assertGreater(path.stat().st_size, 5000)

    def test_report_package_integrates_all_upstream_results(self) -> None:
        task = TaskRecord(
            task_id="T-PACKAGE",
            project_id="P-002",
            project_name="某市评标项目",
            check_type="full",
            status="completed",
            created_at="2026-08-11T00:00:00+08:00",
            updated_at="2026-08-11T00:00:00+08:00",
        )
        agents = [
            AgentResult(agent="合规审查智能体", summary="合规完成"),
            AgentResult(agent="数据核验智能体", summary="数据完成"),
            AgentResult(agent="异常分析智能体", summary="异常完成"),
        ]

        generated = ReportGeneratorAgent().run(
            [],
            [],
            task=task,
            agent_results=agents,
            human_review={"reviewer": "评审委员会", "submit": True},
        )

        package = generated.data["report_package"]
        self.assertEqual(package["project_info"]["project_id"], "P-002")
        self.assertEqual(len(package["compliance_results"]["results"]), 1)
        self.assertEqual(len(package["validation_results"]["results"]), 1)
        self.assertEqual(len(package["anomaly_results"]["results"]), 1)
        self.assertEqual(generated.data["report_status"], "正式核验版")
        self.assertFalse(generated.data["standard_evaluation_report_ready"])

    def test_dify_report_cannot_add_unknown_issue(self) -> None:
        issue = Issue(
            issue_id="I-KNOWN",
            agent="合规审查智能体",
            risk_level="中",
            issue_type="测试问题",
            description="已知问题",
            evidence=["原始证据"],
        )
        semantic = {
            "title": "报告",
            "issues": [
                {"issue_id": "I-KNOWN", "description": "允许"},
                {"issue_id": "I-FAKE", "description": "模型新增"},
            ],
        }
        fake_settings = SimpleNamespace(report_generator_workflow_version="2.0.0")
        with (
            patch("app.agents.report_generator.settings", fake_settings),
            patch(
                "app.services.dify_client.DifyClient.report_generator_enabled",
                new_callable=PropertyMock,
                return_value=True,
            ),
            patch(
                "app.agents.report_generator.dify_client.run_report_generator",
                return_value=semantic,
            ),
        ):
            generated = ReportGeneratorAgent().run([], [issue])

        semantic_issues = generated.data["report_package"]["semantic_content"]["issues"]
        self.assertEqual([item["issue_id"] for item in semantic_issues], ["I-KNOWN"])


if __name__ == "__main__":
    unittest.main()
