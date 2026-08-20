"""Generate a standalone three-part evaluation-rule execution report.

Usage:
    python scripts/generate_three_part_rule_report.py TASK_ID
"""

from __future__ import annotations

import json
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TASK_DIR = ROOT / "data" / "tasks"
REPORT_DIR = ROOT / "data" / "reports"

STATUS_LABELS = {
    "passed": "通过",
    "confirmed_issue": "明确问题",
    "human_review": "待人工复核",
    "insufficient_data": "资料不足未执行",
    "not_applicable": "本次不适用",
    "disabled": "已停用/合并",
}

STATUS_COLORS = {
    "passed": "E8F5E9",
    "confirmed_issue": "FDECEC",
    "human_review": "FFF4E5",
    "insufficient_data": "F2F4F7",
    "not_applicable": "EEF2F6",
    "disabled": "E5E7EB",
}


def _set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def _set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_run_font(run, size=10.5, bold=False, color="202124") -> None:
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _write_cell(cell, value, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=8.5) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    _set_run_font(p.add_run(str(value)), size=size, bold=bold)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Title", 26, "0B2545", 0, 12),
        ("Heading 1", 16, "1F4D78", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 11.5, "1F4D78", 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(header.add_run("三部分评标规则执行报告"), size=8.5, color="667085")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(footer.add_run("招投标全过程智能核验 · 内部工作文件"), size=8, color="667085")


def _find_matrix(task: dict) -> dict:
    result = task.get("result") or {}
    agents = result.get("agent_results") or []
    for agent in reversed(agents):
        data = agent.get("data") if isinstance(agent, dict) else None
        matrix = data.get("three_part_rule_execution") if isinstance(data, dict) else None
        if isinstance(matrix, dict):
            return matrix
    review = task.get("review_request") or {}
    for agent in reversed(review.get("agent_results") or []):
        data = agent.get("data") if isinstance(agent, dict) else None
        matrix = data.get("three_part_rule_execution") if isinstance(data, dict) else None
        if isinstance(matrix, dict):
            return matrix
    raise ValueError("任务中不存在 three_part_rule_execution 数据")


def _add_metadata(doc: Document, rows: list[tuple[str, object]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.0)
    table.columns[1].width = Cm(12.5)
    for label, value in rows:
        cells = table.add_row().cells
        _set_cell_fill(cells[0], "E8EEF5")
        _write_cell(cells[0], label, bold=True, size=9.5)
        _write_cell(cells[1], value, size=9.5)


def _add_summary_table(doc: Document, summary: dict) -> None:
    labels = (
        ("规则总数", "total"), ("实际执行", "executed"), ("通过", "passed"),
        ("明确问题", "confirmed_issue"), ("待人工复核", "human_review"),
        ("资料不足未执行", "insufficient_data"), ("本次不适用", "not_applicable"),
        ("停用/合并", "disabled"),
    )
    table = doc.add_table(rows=2, cols=8)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (label, key) in enumerate(labels):
        _set_cell_fill(table.rows[0].cells[i], "E8EEF5")
        _write_cell(table.rows[0].cells[i], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=8)
        _write_cell(table.rows[1].cells[i], summary.get(key, 0), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)


def _add_group_table(doc: Document, group: dict) -> None:
    rules = group.get("rules") or []
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ("规则编号", "核验类别", "核验事项", "责任智能体", "风险", "执行状态")
    widths = (Cm(1.7), Cm(3.0), Cm(4.4), Cm(3.1), Cm(1.2), Cm(3.3))
    for cell, label, width in zip(table.rows[0].cells, headers, widths, strict=True):
        cell.width = width
        _set_cell_fill(cell, "1F4D78")
        _write_cell(cell, label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=8)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    _set_repeat_table_header(table.rows[0])
    for rule in rules:
        cells = table.add_row().cells
        values = (
            rule.get("rule_id", ""), rule.get("category", ""), rule.get("item", ""),
            rule.get("owner_agent", ""), rule.get("risk_level", ""),
            STATUS_LABELS.get(rule.get("status"), rule.get("status", "")),
        )
        for cell, value, width in zip(cells, values, widths, strict=True):
            cell.width = width
            _write_cell(cell, value, align=WD_ALIGN_PARAGRAPH.CENTER if cell in (cells[0], cells[4], cells[5]) else WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_fill(cells[5], STATUS_COLORS.get(rule.get("status"), "FFFFFF"))


def _add_rule_details(doc: Document, rules: list[dict], statuses: set[str], title: str) -> None:
    selected = [rule for rule in rules if rule.get("status") in statuses]
    doc.add_heading(title, level=2)
    if not selected:
        doc.add_paragraph("无。")
        return
    for rule in selected:
        label = STATUS_LABELS.get(rule.get("status"), rule.get("status", ""))
        doc.add_heading(f"{rule.get('rule_id')}  {rule.get('item')}（{label}）", level=3)
        _add_metadata(doc, [
            ("规则分组", rule.get("group_name", "")),
            ("核验类别", rule.get("category", "")),
            ("责任智能体", rule.get("owner_agent", "")),
            ("执行方式", rule.get("execution_mode", "")),
            ("风险等级", rule.get("risk_level", "")),
            ("执行结论", rule.get("reason", "")),
            ("执行证据", "\n".join(rule.get("execution_evidence") or []) or "未形成可用执行证据"),
            ("计算/比对过程", rule.get("calculation") or "无"),
            ("缺失输入", "\n".join(rule.get("missing_inputs") or []) or "无"),
        ])


def generate(task_id: str) -> Path:
    task_path = TASK_DIR / f"{task_id}.json"
    task = json.loads(task_path.read_text(encoding="utf-8"))
    matrix = _find_matrix(task)
    summary = matrix.get("summary") or {}
    groups = matrix.get("groups") or []
    all_rules = [rule for group in groups for rule in (group.get("rules") or [])]

    doc = Document()
    _configure_document(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(75)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(kicker.add_run("评标报告智能核验"), size=11, bold=True, color="7A5A00")
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("三部分规则执行报告")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(45)
    _set_run_font(subtitle.add_run(task.get("project_name") or "未命名项目"), size=15, bold=True, color="1F4D78")
    _add_metadata(doc, [
        ("任务编号", task_id),
        ("项目编号", task.get("project_id") or "未提供"),
        ("规则版本", matrix.get("ruleset_version") or matrix.get("version") or "未记录"),
        ("报告状态", "待人工复核版" if summary.get("human_review") else "规则执行结果版"),
        ("生成时间", datetime.now().astimezone().strftime("%Y-%m-%d %H:%M:%S %z")),
    ])
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(18)
    _set_run_font(note.add_run("内部工作文件 · 规则执行结果不替代法定评审结论"), size=9, bold=True, color="9B1C1C")
    doc.add_page_break()

    doc.add_heading("一、执行摘要", level=1)
    doc.add_paragraph(
        f"本次按三部分规则体系对任务资料进行核验，共纳入{summary.get('total', 0)}条规则。"
        f"实际具备执行条件并完成核验{summary.get('executed', 0)}条，其中通过{summary.get('passed', 0)}条、"
        f"明确问题{summary.get('confirmed_issue', 0)}条、待人工复核{summary.get('human_review', 0)}条；"
        f"另有{summary.get('insufficient_data', 0)}条因资料不足未执行、"
        f"{summary.get('not_applicable', 0)}条本次不适用、{summary.get('disabled', 0)}条停用或合并。"
    )
    _add_summary_table(doc, summary)
    callout = doc.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    _set_cell_fill(callout.cell(0, 0), "FFF8E8")
    _write_cell(
        callout.cell(0, 0),
        "重要口径：‘资料不足未执行’不等于‘核验通过’；‘待人工复核’不等于‘明确问题’。"
        "人工完成证据核对后，方可将事项更新为明确问题或通过。",
        bold=True,
        size=9.5,
    )

    doc.add_heading("二、三部分规则总体情况", level=1)
    for group in groups:
        counts = Counter(rule.get("status") for rule in group.get("rules") or [])
        doc.add_heading(f"{group.get('group_code')}  {group.get('group_name')}", level=2)
        doc.add_paragraph(
            f"本部分共{group.get('rule_count', len(group.get('rules') or []))}条规则；"
            f"通过{counts['passed']}条、明确问题{counts['confirmed_issue']}条、待人工复核{counts['human_review']}条、"
            f"资料不足未执行{counts['insufficient_data']}条、本次不适用{counts['not_applicable']}条、"
            f"停用/合并{counts['disabled']}条。"
        )
        _add_group_table(doc, group)

    doc.add_heading("三、已执行规则及待人工复核事项", level=1)
    _add_rule_details(doc, all_rules, {"passed", "confirmed_issue", "human_review"}, "3.1 已执行规则明细")

    doc.add_heading("四、资料不足未执行规则", level=1)
    doc.add_paragraph(
        "下列规则因缺少结构化字段、原始证据、评分明细、开标记录、委员会信息或明确判定标准而未执行。"
        "系统未将这些规则记为通过，需补齐对应材料后重新运行。"
    )
    by_group = {}
    for rule in all_rules:
        if rule.get("status") == "insufficient_data":
            by_group.setdefault(rule.get("group_code", ""), []).append(rule)
    for group_code, rules in by_group.items():
        group_name = rules[0].get("group_name", "")
        doc.add_heading(f"{group_code}  {group_name}", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ("规则编号", "核验事项", "缺失输入", "责任智能体")
        for cell, label in zip(table.rows[0].cells, headers, strict=True):
            _set_cell_fill(cell, "E8EEF5")
            _write_cell(cell, label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=8.5)
        _set_repeat_table_header(table.rows[0])
        for rule in rules:
            values = (
                rule.get("rule_id", ""), rule.get("item", ""),
                "；".join(rule.get("missing_inputs") or []) or "未明确",
                rule.get("owner_agent", ""),
            )
            for cell, value in zip(table.add_row().cells, values, strict=True):
                _write_cell(cell, value, size=8.2)

    doc.add_heading("五、本次不适用及停用规则", level=1)
    _add_rule_details(doc, all_rules, {"not_applicable", "disabled"}, "5.1 规则明细")

    doc.add_heading("六、结论与后续处理", level=1)
    if summary.get("human_review"):
        doc.add_paragraph(
            f"本次规则执行未形成自动确认的明确问题，形成{summary.get('human_review')}项待人工复核事项。"
            "建议优先核对相应原始文件、业务系统记录和规则依据，完成后将每项结论更新为“明确问题”或“通过”。"
        )
    else:
        doc.add_paragraph("本次规则执行未形成待人工复核事项，具体结论以各规则明细为准。")
    doc.add_paragraph(
        "对于资料不足未执行的规则，应按缺失输入清单补充采购文件、开标记录、评分明细、评审结果、"
        "委员会组成、签章检测结果等材料后重新执行，不得直接视为核验通过。"
    )
    doc.add_heading("人工复核签署", level=2)
    doc.add_paragraph("复核人：________________    复核日期：________年____月____日")
    doc.add_paragraph("复核结论：____________________________________________________________")

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    output = REPORT_DIR / f"{task_id}_三部分规则执行报告.docx"
    doc.save(output)
    return output


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("Usage: python scripts/generate_three_part_rule_report.py TASK_ID")
    print(generate(sys.argv[1]))
