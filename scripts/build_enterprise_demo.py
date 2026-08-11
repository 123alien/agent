from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "test_data" / "enterprise_demo"
ASSETS = OUT / "assets"
PROJECT_NAME = "XX市信息化平台升级建设项目"
PROJECT_NO = "XXCG-2026-0811"

BIDDERS = [
    {
        "code": "A",
        "name": "华诚科技有限公司",
        "quote": 9_230_000,
        "ip": "120.10.15.21",
        "mac": "AA-01-6F-2C-91-10",
        "machine": "PC-A001",
        "contact": "张伟",
        "phone": "13900001001",
        "email": "zhangwei@huacheng.example",
        "author": "华诚科技投标部",
        "tool": "WPS Office 12.1.0.16388",
        "created": "2026-08-08T09:15:00+08:00",
        "seal_name": "华诚科技有限公司",
    },
    {
        "code": "B",
        "name": "博远信息技术有限公司",
        "quote": 9_600_000,
        "ip": "117.20.33.15",
        "mac": "BC-22-18-7D-90-01",
        "machine": "PC-X888",
        "contact": "李明",
        "phone": "13800001111",
        "email": "bid2026@example.com",
        "author": "投标文件制作中心",
        "tool": "WPS Office 12.1.0.16388",
        "created": "2026-08-08T10:15:00+08:00",
        "seal_name": "博远信息技术有限公司",
    },
    {
        "code": "C",
        "name": "新联科技有限公司",
        "quote": 9_800_000,
        "ip": "117.20.33.15",
        "mac": "BC-22-18-7D-90-01",
        "machine": "PC-X888",
        "contact": "王军",
        "phone": "13800001111",
        "email": "bid2026@example.com",
        "author": "投标文件制作中心",
        "tool": "WPS Office 12.1.0.16388",
        "created": "2026-08-08T10:16:00+08:00",
        "seal_name": "新联科技有限公司",
    },
    {
        "code": "D",
        "name": "天远科技有限公司",
        "quote": 10_000_000,
        "ip": "61.178.20.81",
        "mac": "DD-70-33-2A-61-09",
        "machine": "PC-D204",
        "contact": "陈涛",
        "phone": "13700001004",
        "email": "chentao@tianyuan.example",
        "author": "天远科技商务部",
        "tool": "Microsoft Office 2021",
        "created": "2026-08-08T11:26:00+08:00",
        "seal_name": "天成科技有限公司",
    },
]


def _font_path() -> str:
    for candidate in (
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ):
        if candidate.exists():
            return str(candidate)
    raise FileNotFoundError("未找到可用于测试印章的中文字体")


def make_seal(name: str, output: Path) -> None:
    image = Image.new("RGBA", (640, 640), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)
    red = (205, 20, 35, 235)
    draw.ellipse((34, 34, 606, 606), outline=red, width=18)
    draw.ellipse((72, 72, 568, 568), outline=red, width=5)
    draw.polygon([(320, 205), (348, 288), (436, 288), (365, 339), (392, 425),
                  (320, 374), (248, 425), (275, 339), (204, 288), (292, 288)], fill=red)
    # Keep the company name inside the inner ring. A larger font lets the
    # circle cut through the edge characters and creates avoidable OCR noise.
    title_font = ImageFont.truetype(_font_path(), 38)
    test_font = ImageFont.truetype(_font_path(), 40)
    label = name
    bbox = draw.textbbox((0, 0), label, font=title_font)
    draw.text(((640 - (bbox[2] - bbox[0])) / 2, 105), label, fill=red, font=title_font)
    test_label = "系统测试专用章"
    bbox = draw.textbbox((0, 0), test_label, font=test_font)
    draw.text(((640 - (bbox[2] - bbox[0])) / 2, 500), test_label, fill=red, font=test_font)
    image.save(output)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_doc(doc: Document, title: str, author: str, created: str) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    header = section.header.paragraphs[0]
    header.text = f"{PROJECT_NAME}  |  AI核验测试样本"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(100, 110, 120)
    footer = section.footer.paragraphs[0]
    footer.text = "仅供系统测试，不构成真实投标材料或事实认定"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(130, 130, 130)
    props = doc.core_properties
    props.title = title
    props.author = author
    props.last_modified_by = author
    props.subject = PROJECT_NAME
    props.comments = "AI招投标智能体测试样本"
    props.created = datetime.fromisoformat(created).astimezone(timezone.utc).replace(tzinfo=None)
    props.modified = props.created


def add_cover(doc: Document, bidder: dict | None = None) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(PROJECT_NAME)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(31, 77, 120)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("采购文件（测试样本）" if bidder is None else "投标响应文件（测试样本）")
    r.bold = True
    r.font.size = Pt(20)
    doc.add_paragraph()
    table = doc.add_table(rows=4 if bidder else 3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    rows = [("项目编号", PROJECT_NO), ("项目名称", PROJECT_NAME)]
    if bidder:
        rows.extend([("投标人", bidder["name"]), ("投标报价", f"人民币{bidder['quote'] / 10000:.0f}万元")])
    else:
        rows.append(("招标控制价", "人民币1000万元"))
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].width = Inches(1.875)
        row.cells[1].width = Inches(4.625)
        set_cell_shading(row.cells[0], "F2F4F7")
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph("特别说明：本文件为基于公开典型风险特征构造的虚拟测试材料，企业、人员、地址及联系方式均为虚构。")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(155, 28, 28)
        run.bold = True


def add_page(doc: Document, number: int, title: str, paragraphs: list[str], table_rows=None, seal=None) -> None:
    if number > 1:
        doc.add_page_break()
    p = doc.add_paragraph(f"第{number}页  {title}", style="Heading 1")
    p.paragraph_format.keep_with_next = True
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for row_index, values in enumerate(table_rows):
            for column_index, value in enumerate(values):
                table.cell(row_index, column_index).text = str(value)
                if row_index == 0:
                    set_cell_shading(table.cell(row_index, column_index), "E8EEF5")
                    for run in table.cell(row_index, column_index).paragraphs[0].runs:
                        run.bold = True
    if seal:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run().add_picture(str(seal), width=Cm(4.2))


def bidder_pages(bidder: dict) -> list[tuple[str, list[str], list[list[str]] | None, str | None]]:
    shared_1 = "本项目采用微服务架构，提高系统扩展性和运行稳定性，建设统一服务治理、统一日志和统一监控能力。"
    shared_2 = "项目实施过程中采用双周迭代机制，由项目经理统一协调资源并形成阶段性交付成果。"
    shared_typo = "系统可用性达到99.99%%，系统应具备统一身份正认功能，实现用户、角色和权限的统一管理。"
    tech = [
        "我公司已充分理解采购需求，将遵循安全、可靠、可维护的建设原则完成平台升级。",
        "采用分层架构与标准接口，支持业务组件解耦、横向扩展和统一运维。",
    ]
    if bidder["code"] in {"B", "C"}:
        tech = [shared_1, shared_2, shared_typo]
        if bidder["code"] == "C":
            tech.append("在不改变总体技术路线的前提下，按阶段完成联调、试运行与验收交付。")
    pages = [
        ("封面", [], None, None),
        ("投标函", [f"致：XX市政务服务管理中心。我方愿以人民币{bidder['quote'] / 10000:.0f}万元承担本项目全部工作。", "我方承诺投标有效期为90日，并对响应文件真实性负责。"], None, None),
        ("开标一览表", ["报价包含软件开发、系统集成、数据迁移、培训和两年运维服务。"], [["投标人", "投标报价（元）", "工期", "质保期"], [bidder["name"], bidder["quote"], "180日历天", "24个月"]], None),
        ("法定代表人身份证明", [f"兹证明{bidder['contact']}同志为我公司法定代表人，有权代表本公司参加本项目投标活动。"], None, None),
        ("授权委托书", [f"现委托{bidder['contact']}为本项目授权代表，联系电话：{bidder['phone']}，电子邮箱：{bidder['email']}。"], None, None),
        ("资格证明材料", ["我公司具备独立承担民事责任的能力，具有良好商业信誉和履行合同所必需的专业技术能力。", "本页证照编号、住所与有效期限均为虚构测试数据。"], [["证照", "编号", "状态"], ["营业执照", f"TEST-{bidder['code']}-91310000", "有效"], ["信息安全管理体系", f"ISMS-{bidder['code']}-2026", "有效"]], None),
        ("类似项目业绩", ["以下业绩仅为模拟数据，不对应现实合同。"], [["项目", "合同金额", "验收日期"], ["政务平台运维服务", "680万元", "2025-12-18"], ["数据共享平台升级", "520万元", "2024-11-03"]], None),
        ("投标函附录及盖章页", ["投标人（盖章）：", f"投标人名称：{bidder['name']}", "日期：2026年8月9日"], None, "seal8"),
        ("技术方案总体设计", tech, None, None),
        ("项目实施计划", ["项目分为启动、需求确认、开发配置、联调测试、试运行和验收六个阶段。", "建立周例会、问题台账、变更控制和质量门禁机制。"], [["阶段", "计划周期", "主要交付物"], ["启动与需求确认", "第1—3周", "项目章程、需求规格"], ["开发与联调", "第4—18周", "系统版本、测试记录"], ["试运行与验收", "第19—26周", "试运行报告、验收材料"]], None),
        ("运维与服务保障", ["服务保障以服务响应时间、服务方案、技术支持能力和履约保障能力为核心。", "重大故障15分钟响应、2小时内提出处置方案，重要操作形成审计记录。"], None, None),
        ("商务偏离表及盖章页", ["投标人（盖章）：", f"投标人名称：{bidder['name']}", "我公司对采购文件商务条款无负偏离。"], [["条款", "响应情况", "偏离说明"], ["工期", "完全响应", "无"], ["质保期", "完全响应", "无"]], "seal12"),
        ("技术偏离表", ["我公司承诺满足采购文件所列功能、性能、接口、安全和交付要求。"], [["序号", "技术要求", "响应"], [1, "统一身份认证", "完全响应"], [2, "日志审计", "完全响应"], [3, "数据交换", "完全响应"]], None),
        ("报价明细表", ["各分项价格合计等于投标总价。"], [["分项", "金额（元）"], ["软件开发", int(bidder["quote"] * 0.52)], ["系统集成", int(bidder["quote"] * 0.28)], ["运维服务", bidder["quote"] - int(bidder["quote"] * 0.52) - int(bidder["quote"] * 0.28)], ["合计", bidder["quote"]]], None),
        ("签署页", ["法定代表人（签字）：", f"投标人：{bidder['name']}", "日期：2026年8月9日"], None, None),
    ]
    return pages


def build_bidder_docx(bidder: dict) -> Path:
    doc = Document()
    configure_doc(doc, f"{bidder['name']}投标响应文件", bidder["author"], bidder["created"])
    seals = {}
    for page in (8, 12):
        if bidder["code"] == "D" and page == 8:
            continue
        seal_path = ASSETS / f"seal_{bidder['code']}_{page}.png"
        make_seal(bidder["seal_name"] if page == 12 else bidder["name"], seal_path)
        seals[f"seal{page}"] = seal_path
    for page_number, (title, paras, rows, seal_key) in enumerate(bidder_pages(bidder), start=1):
        if page_number == 1:
            add_cover(doc, bidder)
        else:
            add_page(doc, page_number, title, paras, rows, seals.get(seal_key) if seal_key else None)
    path = OUT / f"{bidder['code']}_{bidder['name']}_投标响应文件.docx"
    doc.save(path)
    return path


def build_tender_docx() -> Path:
    doc = Document()
    configure_doc(doc, "采购文件", "XX市公共资源交易中心", "2026-08-01T09:00:00+08:00")
    add_cover(doc)
    tender_pages = [
        ("投标人须知", ["采购人：XX市政务服务管理中心。采购代理机构：XX市公共资源交易中心。", "投标截止时间：2026年8月10日09时30分；开标地点：XX市公共资源交易中心第三开标室。"]),
        ("项目概况与采购范围", ["建设内容包括政务信息化平台升级、数据迁移、系统集成、培训及两年运维服务。", "招标控制价：人民币1000万元。超过最高投标限价的投标将被否决。"]),
        ("投标人资格要求", ["投标人应具备独立承担民事责任的能力。", "投标人注册地址必须位于本市，注册地址位于本市以外的投标人不得参加本项目投标。", "投标人须提供2023年1月1日至投标截止日期间承担的政务信息化项目业绩。"]),
        ("技术要求", ["平台应采用开放接口并满足网络安全、数据安全和日志审计要求。", "监控工具必须使用甲公司生产的“城市云眼V5”产品，不接受其他品牌或具有同等功能的替代产品。"]),
        ("评标办法", ["本项目采用综合评分法，总分100分。"], [["评审因素", "分值"], ["投标报价", 30], ["技术方案", 40], ["项目实施能力", 20], ["服务方案", 10]]),
        ("投标文件格式", ["投标文件应包含投标函、报价表、资格证明、技术方案、商务及技术偏离表和签署页。", "标注“投标人（盖章）”处应加盖公章，标注“法定代表人（签字）”处应由本人签署。"]),
        ("合同主要条款", ["合同工期为180日历天，质保期24个月。", "付款、验收、知识产权、保密及违约责任以最终合同为准。"]),
    ]
    for index, (title, paras, *maybe_table) in enumerate(tender_pages, start=2):
        add_page(doc, index, title, paras, maybe_table[0] if maybe_table else None)
    path = OUT / "00_采购文件_XX市信息化平台升级建设项目.docx"
    doc.save(path)
    return path


def pdf_styles():
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    return {
        "title": ParagraphStyle("title-cn", fontName="STSong-Light", fontSize=22, leading=30, alignment=TA_CENTER, textColor=colors.HexColor("#1F4D78"), spaceAfter=20),
        "h1": ParagraphStyle("h1-cn", fontName="STSong-Light", fontSize=16, leading=22, textColor=colors.HexColor("#2E74B5"), spaceAfter=12),
        "body": ParagraphStyle("body-cn", fontName="STSong-Light", fontSize=10.5, leading=18, alignment=TA_JUSTIFY, spaceAfter=8),
        "note": ParagraphStyle("note-cn", fontName="STSong-Light", fontSize=9, leading=14, alignment=TA_CENTER, textColor=colors.HexColor("#9B1C1C")),
    }


def pdf_page(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("STSong-Light", 8)
    canvas.setFillColor(colors.HexColor("#6B7280"))
    canvas.drawString(2.2 * cm, 1.4 * cm, "仅供AI招投标智能体测试，不构成真实事实认定")
    canvas.drawRightString(18.8 * cm, 1.4 * cm, f"第 {doc.page} 页")
    canvas.restoreState()


def add_pdf_page(story, styles, number, title, paras, rows=None, seal=None, cover_bidder=None):
    if number > 1:
        story.append(PageBreak())
    if number == 1:
        story.extend([Spacer(1, 5 * cm), Paragraph(PROJECT_NAME, styles["title"]), Paragraph("采购文件（测试样本）" if cover_bidder is None else "投标响应文件（测试样本）", styles["title"])])
        meta = [["项目编号", PROJECT_NO], ["项目名称", PROJECT_NAME]]
        if cover_bidder:
            meta.extend([["投标人", cover_bidder["name"]], ["投标报价", f"人民币{cover_bidder['quote']/10000:.0f}万元"]])
        else:
            meta.append(["招标控制价", "人民币1000万元"])
        table = Table(meta, colWidths=[4.2 * cm, 11.8 * cm])
        table.setStyle(TableStyle([("FONTNAME", (0, 0), (-1, -1), "STSong-Light"), ("GRID", (0, 0), (-1, -1), .5, colors.HexColor("#B8C2CC")), ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F4F7")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8)]))
        story.extend([table, Spacer(1, .5 * cm), Paragraph("特别说明：本文件全部主体和数据均为虚构，仅供系统测试。", styles["note"])])
        return
    story.append(Paragraph(f"第{number}页  {title}", styles["h1"]))
    for text in paras:
        story.append(Paragraph(text, styles["body"]))
    if rows:
        table = Table(rows, repeatRows=1)
        table.setStyle(TableStyle([("FONTNAME", (0, 0), (-1, -1), "STSong-Light"), ("FONTSIZE", (0, 0), (-1, -1), 8.5), ("GRID", (0, 0), (-1, -1), .5, colors.HexColor("#9CA3AF")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6)]))
        story.extend([Spacer(1, .2 * cm), table])
    if seal:
        img = RLImage(str(seal), width=4.2 * cm, height=4.2 * cm)
        img.hAlign = "RIGHT"
        story.extend([Spacer(1, .5 * cm), img])


def build_bidder_pdf(bidder: dict) -> Path:
    path = OUT / f"{bidder['code']}_{bidder['name']}_投标响应文件.pdf"
    styles = pdf_styles()
    seals = {}
    for page in (8, 12):
        seal_path = ASSETS / f"seal_{bidder['code']}_{page}.png"
        if seal_path.exists():
            seals[f"seal{page}"] = seal_path
    story = []
    for page_number, (title, paras, rows, seal_key) in enumerate(bidder_pages(bidder), start=1):
        add_pdf_page(story, styles, page_number, title, paras, rows, seals.get(seal_key) if seal_key else None, bidder)
    SimpleDocTemplate(str(path), pagesize=A4, leftMargin=2.54 * cm, rightMargin=2.54 * cm, topMargin=2.3 * cm, bottomMargin=2.2 * cm, title=path.stem, author=bidder["author"]).build(story, onFirstPage=pdf_page, onLaterPages=pdf_page)
    return path


def build_tender_pdf() -> Path:
    path = OUT / "00_采购文件_XX市信息化平台升级建设项目.pdf"
    styles = pdf_styles()
    docx = build_tender_docx()
    del docx
    tender_pages = [
        ("封面", [], None),
        ("投标人须知", ["采购人：XX市政务服务管理中心。采购代理机构：XX市公共资源交易中心。", "投标截止时间：2026年8月10日09时30分；开标地点：XX市公共资源交易中心第三开标室。"], None),
        ("项目概况与采购范围", ["建设内容包括政务信息化平台升级、数据迁移、系统集成、培训及两年运维服务。", "招标控制价：人民币1000万元。超过最高投标限价的投标将被否决。"], None),
        ("投标人资格要求", ["投标人应具备独立承担民事责任的能力。", "投标人注册地址必须位于本市，注册地址位于本市以外的投标人不得参加本项目投标。", "投标人须提供2023年1月1日至投标截止日期间承担的政务信息化项目业绩。"], None),
        ("技术要求", ["平台应采用开放接口并满足网络安全、数据安全和日志审计要求。", "监控工具必须使用甲公司生产的“城市云眼V5”产品，不接受其他品牌或具有同等功能的替代产品。"], None),
        ("评标办法", ["本项目采用综合评分法，总分100分。"], [["评审因素", "分值"], ["投标报价", 30], ["技术方案", 40], ["项目实施能力", 20], ["服务方案", 10]]),
        ("投标文件格式", ["标注“投标人（盖章）”处应加盖公章，标注“法定代表人（签字）”处应由本人签署。"], None),
        ("合同主要条款", ["合同工期为180日历天，质保期24个月。"], None),
    ]
    story = []
    for page_number, (title, paras, rows) in enumerate(tender_pages, start=1):
        add_pdf_page(story, styles, page_number, title, paras, rows)
    SimpleDocTemplate(str(path), pagesize=A4, leftMargin=2.54 * cm, rightMargin=2.54 * cm, topMargin=2.3 * cm, bottomMargin=2.2 * cm, title=path.stem, author="XX市公共资源交易中心").build(story, onFirstPage=pdf_page, onLaterPages=pdf_page)
    return path


def build_metadata_xlsx() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "项目基准"
    ws.append(["字段", "值", "说明"])
    ws.append(["project_id", PROJECT_NO, "虚拟测试项目"])
    ws.append(["project_name", PROJECT_NAME, "虚拟测试项目"])
    ws.append(["control_price", 10_000_000, "招标控制价（元）"])
    ws.append(["bid_deadline", "2026-08-10 09:30:00", "投标截止时间"])
    ws.append(["tenderer", "XX市政务服务管理中心", "采购人"])
    bid = wb.create_sheet("投标记录")
    bid.append(["supplier_code", "supplier_name", "bid_price", "upload_time", "contact", "phone", "email"])
    for index, item in enumerate(BIDDERS):
        bid.append([f"S00{index + 1}", item["name"], item["quote"], f"2026-08-10 08:{12 + index * 3}:00", item["contact"], item["phone"], item["email"]])
    meta = wb.create_sheet("文件与网络元数据")
    meta.append(["supplier_code", "supplier_name", "file_author", "created_time", "creation_tool", "upload_ip", "mac_address", "machine_code", "cost_software_lock_id"])
    for index, item in enumerate(BIDDERS):
        lock_id = "LOCK-2026-7788" if item["code"] in {"B", "C"} else f"LOCK-{item['code']}-2026"
        meta.append([f"S00{index + 1}", item["name"], item["author"], item["created"], item["tool"], item["ip"], item["mac"], item["machine"], lock_id])
    score = wb.create_sheet("专家评分")
    score.append(["supplier_code", "supplier_name", "专家一", "专家二", "专家三", "平均分", "排名"])
    rows = [("S001", BIDDERS[0]["name"], 91, 92, 90), ("S002", BIDDERS[1]["name"], 86, 87, 85), ("S003", BIDDERS[2]["name"], 84, 48, 86), ("S004", BIDDERS[3]["name"], 78, 80, 79)]
    for rank, values in enumerate(rows, start=1):
        row = score.max_row + 1
        score.append([*values, f"=AVERAGE(C{row}:E{row})", rank])
    expected = wb.create_sheet("预期风险线索")
    expected.append(["finding_id", "对象", "风险线索", "预期状态", "说明"])
    expected_rows = [
        ("F-001", "B/C", "IP、MAC、机器码、加密锁号重合", "人工复核", "多项独立信号组合，不直接认定串通投标"),
        ("F-002", "B/C", "非模板文本高度相似且含共同错字", "人工复核", "与设备及网络线索组合评估"),
        ("F-003", "B/C/D", "报价960、980、1000万元呈等差数列", "人工复核", "单一报价规律不足以直接定性"),
        ("F-004", "D", "第8页应盖章位置未检测到印章", "人工复核", "status=not_detected"),
        ("F-005", "D", "第12页印章主体与投标人名称不一致", "人工复核", "status=mismatch或low_confidence"),
        ("F-006", "D", "第15页签名模型未配置", "人工复核", "status=not_checked，不得表述为未签字"),
    ]
    for row in expected_rows:
        expected.append(row)
    thin = Side(style="thin", color="C7CDD4")
    for sheet in wb.worksheets:
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = sheet.dimensions
        for cell in sheet[1]:
            cell.fill = PatternFill("solid", fgColor="1F4D78")
            cell.font = Font(color="FFFFFF", bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center")
        for row in sheet.iter_rows():
            for cell in row:
                cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        for col in range(1, sheet.max_column + 1):
            width = max(len(str(sheet.cell(row, col).value or "")) for row in range(1, min(sheet.max_row, 30) + 1))
            sheet.column_dimensions[get_column_letter(col)].width = min(max(width * 1.4, 12), 45)
    path = OUT / "05_电子交易与评审元数据.xlsx"
    wb.save(path)
    return path


def build_expected_json() -> Path:
    payload = {
        "dataset_version": "1.0.0",
        "project": {"project_id": PROJECT_NO, "project_name": PROJECT_NAME, "control_price": 10_000_000},
        "principle": "发现线索、给出风险等级和证据并转人工复核，不直接认定串通投标。",
        "expected_findings": [
            {"id": "F-001", "agents": ["数据核验智能体", "异常分析智能体"], "entities": ["S002", "S003"], "signals": ["upload_ip", "mac_address", "machine_code", "cost_software_lock_id"], "risk_level": "高", "requires_human_review": True, "expected_wording": "检测到多项疑似串通投标风险特征，建议人工进一步核验。"},
            {"id": "F-002", "agents": ["文档解析智能体", "异常分析智能体"], "entities": ["S002", "S003"], "signals": ["文本高度相似", "共同错字99.99%%", "共同错字统一身份正认"], "risk_level": "高", "requires_human_review": True},
            {"id": "F-003", "agents": ["数据核验智能体", "异常分析智能体"], "entities": ["S002", "S003", "S004"], "signals": ["9600000", "9800000", "10000000"], "risk_level": "中", "requires_human_review": True},
            {"id": "F-004", "agents": ["文档解析智能体"], "entities": ["S004"], "page": 8, "expected_status": "not_detected", "requires_human_review": True},
            {"id": "F-005", "agents": ["文档解析智能体"], "entities": ["S004"], "page": 12, "expected_status": ["mismatch", "low_confidence"], "requires_human_review": True},
            {"id": "F-006", "agents": ["文档解析智能体"], "entities": ["S004"], "page": 15, "expected_status": "not_checked", "forbidden_wording": "未签字", "requires_human_review": True},
            {"id": "F-007", "agents": ["合规审查智能体"], "entities": ["采购文件"], "signals": ["注册地址必须位于本市", "指定城市云眼V5且不接受同等产品"], "requires_human_review": True},
        ],
    }
    path = OUT / "expected_findings.json"
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    build_tender_pdf()
    for bidder in BIDDERS:
        build_bidder_docx(bidder)
        build_bidder_pdf(bidder)
    build_metadata_xlsx()
    build_expected_json()
    print(f"Enterprise demo generated: {OUT}")


if __name__ == "__main__":
    main()
