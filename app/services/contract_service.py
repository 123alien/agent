from datetime import date
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.core.config import ensure_data_dirs, settings
from app.schemas.contract import (
    ContractGenerationRequest,
    ContractValidationItem,
)


_CN_DIGITS = "零壹贰叁肆伍陆柒捌玖"
_CN_UNITS = ("", "拾", "佰", "仟")
_CN_GROUP_UNITS = ("", "万", "亿", "兆")


def _set_east_asia_font(run, font_name: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        from docx.oxml import OxmlElement

        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _integer_to_chinese(value: int) -> str:
    if value == 0:
        return "零"
    groups: list[int] = []
    while value:
        groups.append(value % 10000)
        value //= 10000
    parts: list[str] = []
    zero_pending = False
    for group_index in range(len(groups) - 1, -1, -1):
        group = groups[group_index]
        if group == 0:
            zero_pending = bool(parts)
            continue
        if zero_pending or (parts and group < 1000):
            if not parts[-1].endswith("零"):
                parts.append("零")
        zero_pending = False
        group_parts: list[str] = []
        inner_zero = False
        for unit_index in range(3, -1, -1):
            divisor = 10**unit_index
            digit = group // divisor % 10
            if digit:
                if inner_zero and group_parts:
                    group_parts.append("零")
                group_parts.append(_CN_DIGITS[digit] + _CN_UNITS[unit_index])
                inner_zero = False
            elif group_parts:
                inner_zero = True
        parts.append("".join(group_parts) + _CN_GROUP_UNITS[group_index])
    return "".join(parts).rstrip("零")


def amount_to_chinese(amount: Decimal) -> str:
    normalized = amount.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    integer = int(normalized)
    fraction = int((normalized - integer) * 100)
    jiao, fen = divmod(fraction, 10)
    suffix = "整" if fraction == 0 else (
        (_CN_DIGITS[jiao] + "角" if jiao else "零")
        + (_CN_DIGITS[fen] + "分" if fen else "")
    )
    return f"人民币{_integer_to_chinese(integer)}元{suffix}"


def validate_contract(request: ContractGenerationRequest) -> list[ContractValidationItem]:
    items: list[ContractValidationItem] = []
    if request.service_end_date <= request.service_start_date:
        items.append(
            ContractValidationItem(
                code="INVALID_SERVICE_PERIOD",
                level="error",
                message="服务结束日期必须晚于开始日期。",
                requires_human_review=True,
            )
        )
    for role, party in (("采购人", request.purchaser), ("供应商", request.supplier)):
        if not party.unified_social_credit_code:
            items.append(
                ContractValidationItem(
                    code=f"MISSING_{'PURCHASER' if role == '采购人' else 'SUPPLIER'}_CREDIT_CODE",
                    level="warning",
                    message=f"{role}统一社会信用代码缺失。",
                    requires_human_review=True,
                )
            )
        if not party.address:
            items.append(
                ContractValidationItem(
                    code=f"MISSING_{'PURCHASER' if role == '采购人' else 'SUPPLIER'}_ADDRESS",
                    level="warning",
                    message=f"{role}地址缺失。",
                    requires_human_review=True,
                )
            )
    if not request.breach_terms:
        items.append(
            ContractValidationItem(
                code="MISSING_BREACH_TERMS",
                level="warning",
                message="违约责任条款尚未提供，已使用待补充提示。",
                requires_human_review=True,
            )
        )
    items.append(
        ContractValidationItem(
            code="LEGAL_REVIEW_REQUIRED",
            level="info",
            message="合同定稿前必须由业务及法务人员复核。",
            requires_human_review=True,
        )
    )
    return items


def _configure(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.05)
    section.right_margin = Inches(1.05)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)


def _add_clause(document: Document, title: str, paragraphs: list[str]) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(5)
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    _set_east_asia_font(run, "黑体")
    for index, text in enumerate(paragraphs, start=1):
        document.add_paragraph(f"{index}. {text}")


def create_contract_docx(
    contract_id: str,
    contract_number: str,
    request: ContractGenerationRequest,
    validation_items: list[ContractValidationItem],
    supplementary_clauses: dict[str, list[str]] | None = None,
) -> Path:
    ensure_data_dirs()
    path = settings.contracts_dir / f"{contract_id}.docx"
    document = Document()
    _configure(document)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run(request.template_type)
    run.bold = True
    run.font.size = Pt(22)
    _set_east_asia_font(run, "黑体")

    for label, value in (
        ("合同编号", contract_number),
        ("项目编号", request.project_id),
        ("项目名称", request.project_name),
        ("采购人（甲方）", request.purchaser.name),
        ("供应商（乙方）", request.supplier.name),
    ):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}：").bold = True
        paragraph.add_run(value)

    _add_clause(document, "第一条 合同标的与服务范围", request.service_scope)
    _add_clause(
        document,
        "第二条 合同金额",
        [
            f"合同总金额（含税）为人民币￥{request.contract_amount:,.2f}元（大写：{amount_to_chinese(request.contract_amount)}）。",
            "上述价款包含完成本合同约定工作所需的全部费用。",
        ],
    )
    _add_clause(
        document,
        "第三条 服务期限",
        [f"服务期限自{request.service_start_date.isoformat()}起至{request.service_end_date.isoformat()}止。"],
    )
    _add_clause(document, "第四条 付款方式", request.payment_terms)
    _add_clause(document, "第五条 验收标准", request.acceptance_criteria)
    supplementary_clauses = supplementary_clauses or {}
    _add_clause(
        document,
        "第六条 服务水平与考核",
        supplementary_clauses.get("service_level_terms")
        or ["乙方应按照采购文件、响应文件及双方确认的服务水平要求提供服务。"],
    )
    _add_clause(
        document,
        "第七条 数据安全与保密",
        supplementary_clauses.get("data_security_terms")
        or ["乙方应遵守适用的数据安全、网络安全和保密要求，未经甲方书面同意不得向第三方披露项目数据。"],
    )
    _add_clause(
        document,
        "第八条 知识产权",
        supplementary_clauses.get("intellectual_property_terms")
        or ["双方已有知识产权归原权利人所有；项目新增成果的权属及使用范围由双方依据采购文件和响应文件确认。"],
    )
    _add_clause(
        document,
        "第九条 变更管理",
        supplementary_clauses.get("change_management_terms")
        or ["服务范围、进度或费用发生变更的，应履行书面确认程序，未经确认的变更不作为结算依据。"],
    )
    _add_clause(
        document,
        "第十条 违约责任",
        request.breach_terms or ["【待双方结合采购文件、中标文件及项目实际情况补充】"],
    )
    _add_clause(
        document,
        "第十一条 合同解除与终止",
        supplementary_clauses.get("termination_terms")
        or ["出现法定或约定解除情形时，守约方可依法解除合同；合同终止后乙方仍应完成资料、账号和数据移交。"],
    )
    _add_clause(
        document,
        "第十二条 不可抗力",
        supplementary_clauses.get("force_majeure_terms")
        or ["受不可抗力影响的一方应及时通知对方并提供证明，双方根据影响程度协商处理。"],
    )
    _add_clause(document, "第十三条 争议解决", [request.dispute_resolution])
    _add_clause(
        document,
        "第十四条 生效、文件组成及效力顺序",
        [
            "本合同经双方签字并盖章后生效。",
            "本合同、补充协议、中标通知书、采购文件、响应文件及双方确认的其他文件均为合同组成部分；如内容不一致，应结合文件形成时间、约定内容和适用规则进行人工确认。",
        ],
    )

    if validation_items:
        document.add_page_break()
        _add_clause(
            document,
            "生成校验与人工复核清单（签署前删除本页）",
            [item.message for item in validation_items],
        )

    document.add_paragraph("甲方（盖章）：____________________")
    document.add_paragraph("乙方（盖章）：____________________")
    document.add_paragraph(f"签署日期：________年____月____日")
    document.save(path)
    return path
