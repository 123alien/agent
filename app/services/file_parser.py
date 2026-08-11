from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable


@dataclass
class ParsedPage:
    number: int
    text: str


@dataclass
class ParsedTableData:
    page: int | None
    page_end: int | None = None
    sheet: str = ""
    start_row: int | None = None
    continued: bool = False
    rows: list[list[str]] = field(default_factory=list)


@dataclass
class ParsedLayoutElement:
    element_type: str
    text: str = ""
    page: int | None = None
    order: int = 0
    bbox: list[float] = field(default_factory=list)
    source_name: str = ""


@dataclass
class ParsedFileContent:
    text: str
    pages: list[ParsedPage] = field(default_factory=list)
    tables: list[ParsedTableData] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    is_scanned: bool = False
    selected_tool: str = ""
    tool_trace: list[str] = field(default_factory=list)
    ocr_applied: bool = False
    ocr_confidence: float = 0.0
    sheet_names: list[str] = field(default_factory=list)
    layout_elements: list[ParsedLayoutElement] = field(default_factory=list)

    @property
    def page_count(self) -> int:
        return len(self.pages)


def _clean_text(text: str) -> str:
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    cleaned: list[str] = []
    blank = False
    for line in lines:
        if line:
            cleaned.append(line)
            blank = False
        elif cleaned and not blank:
            cleaned.append("")
            blank = True
    return "\n".join(cleaned).strip()


def _read_text_file(file_path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return file_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return file_path.read_text(encoding="utf-8", errors="replace")


def _normalize_table(table: list[list[object | None]]) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table:
        cells = [_clean_text(str(cell or "")) for cell in row]
        if any(cells):
            rows.append(cells)
    return rows


def _parse_text(file_path: Path) -> ParsedFileContent:
    text = _clean_text(_read_text_file(file_path))
    return ParsedFileContent(
        text=text,
        pages=[ParsedPage(number=1, text=text)],
        warnings=[] if text else ["文件未提取到有效文本"],
        layout_elements=[ParsedLayoutElement(element_type="paragraph", text=text, page=1, order=1)] if text else [],
    )


def _parse_pdf(file_path: Path) -> ParsedFileContent:
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("解析 PDF 需要安装 pdfplumber") from exc

    pages: list[ParsedPage] = []
    tables: list[ParsedTableData] = []
    warnings: list[str] = []
    layout_elements: list[ParsedLayoutElement] = []
    with pdfplumber.open(file_path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            page_text = _clean_text(page.extract_text() or "")
            pages.append(ParsedPage(number=page_number, text=page_text))
            words = page.extract_words() or []
            line_groups: list[list[dict]] = []
            for word in words:
                if not line_groups or abs(float(word["top"]) - float(line_groups[-1][0]["top"])) > 4:
                    line_groups.append([word])
                else:
                    line_groups[-1].append(word)
            for order, group in enumerate(line_groups, start=1):
                line_text = " ".join(str(word["text"]) for word in group).strip()
                if not line_text:
                    continue
                bbox = [min(float(w["x0"]) for w in group), min(float(w["top"]) for w in group), max(float(w["x1"]) for w in group), max(float(w["bottom"]) for w in group)]
                element_type = "title" if len(line_text) <= 60 and any(pattern.match(line_text) for pattern in (re.compile(r"^第.+[章节]"), re.compile(r"^[一二三四五六七八九十]+、"))) else "paragraph"
                layout_elements.append(ParsedLayoutElement(element_type=element_type, text=line_text, page=page_number, order=order, bbox=bbox))
            try:
                for table_object in page.find_tables() or []:
                    rows = _normalize_table(table_object.extract())
                    if rows:
                        tables.append(ParsedTableData(page=page_number, page_end=page_number, rows=rows))
                        layout_elements.append(ParsedLayoutElement(element_type="table", text="\n".join(" | ".join(row) for row in rows), page=page_number, order=len(line_groups) + len(tables), bbox=[float(v) for v in table_object.bbox]))
            except Exception:
                warnings.append(f"第 {page_number} 页表格提取失败")

    text = "\n\n".join(page.text for page in pages if page.text)
    nonempty_pages = sum(bool(page.text) for page in pages)
    is_scanned = bool(pages) and (
        nonempty_pages == 0 or len(text) / max(len(pages), 1) < 20
    )
    if is_scanned:
        warnings.append("PDF 文本密度过低，可能是扫描件，建议进行 OCR 或人工复核")
    if not text:
        warnings.append("PDF 未提取到有效文本")
    return ParsedFileContent(
        text=text,
        pages=pages,
        tables=_merge_continued_tables(tables),
        layout_elements=layout_elements,
        warnings=warnings,
        is_scanned=is_scanned,
    )


def _parse_docx(file_path: Path) -> ParsedFileContent:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("解析 Word 需要安装 python-docx") from exc

    doc = Document(file_path)
    parts: list[str] = []
    layout_elements: list[ParsedLayoutElement] = []
    for paragraph in doc.paragraphs:
        text = _clean_text(paragraph.text)
        if text:
            parts.append(text)
            style_name = getattr(paragraph.style, "name", "") or ""
            layout_elements.append(ParsedLayoutElement(element_type="title" if style_name.lower().startswith("heading") or style_name.startswith("标题") else "paragraph", text=text, page=1, order=len(layout_elements) + 1))

    tables: list[ParsedTableData] = []
    for table in doc.tables:
        rows = [[_clean_text(cell.text) for cell in row.cells] for row in table.rows]
        rows = [row for row in rows if any(row)]
        if rows:
            tables.append(ParsedTableData(page=None, rows=rows))
            parts.extend(" | ".join(row) for row in rows)
            layout_elements.append(ParsedLayoutElement(element_type="table", text="\n".join(" | ".join(row) for row in rows), page=1, order=len(layout_elements) + 1))

    text = _clean_text("\n".join(parts))
    warnings = [] if text else ["Word 文档未提取到有效文本"]
    return ParsedFileContent(
        text=text,
        pages=[ParsedPage(number=1, text=text)],
        tables=tables,
        warnings=warnings,
        layout_elements=layout_elements,
    )


def _parse_xlsx(file_path: Path) -> ParsedFileContent:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("解析 XLSX 需要安装 openpyxl") from exc

    workbook = load_workbook(file_path, data_only=True, read_only=True)
    tables: list[ParsedTableData] = []
    parts: list[str] = []
    warnings: list[str] = []
    for sheet in workbook.worksheets:
        rows: list[list[str]] = []
        for row in sheet.iter_rows(values_only=True):
            values = [_clean_text(str(value)) if value is not None else "" for value in row]
            while values and not values[-1]:
                values.pop()
            if any(values):
                rows.append(values)
        if rows:
            tables.append(ParsedTableData(page=None, sheet=sheet.title, start_row=1, rows=rows))
            parts.append(f"【工作表：{sheet.title}】")
            parts.extend(" | ".join(row) for row in rows)
        else:
            warnings.append(f"工作表“{sheet.title}”未读取到数据")
    sheet_names = list(workbook.sheetnames)
    workbook.close()
    text = _clean_text("\n".join(parts))
    if not text:
        warnings.append("Excel 未提取到有效数据")
    return ParsedFileContent(
        text=text,
        pages=[ParsedPage(number=1, text=text)],
        tables=tables,
        warnings=warnings,
        sheet_names=sheet_names,
        layout_elements=[ParsedLayoutElement(element_type="table", text="\n".join(" | ".join(row) for row in table.rows), page=None, order=index, source_name=table.sheet) for index, table in enumerate(tables, start=1)],
    )


def _normalized_header(row: list[str]) -> tuple[str, ...]:
    return tuple(re.sub(r"\s+", "", cell).lower() for cell in row)


def _merge_continued_tables(tables: list[ParsedTableData]) -> list[ParsedTableData]:
    merged: list[ParsedTableData] = []
    for table in tables:
        if not merged or table.page is None or merged[-1].page_end is None:
            merged.append(table)
            continue
        previous = merged[-1]
        consecutive = table.page == previous.page_end + 1
        same_columns = bool(previous.rows and table.rows and len(previous.rows[0]) == len(table.rows[0]))
        same_header = same_columns and _normalized_header(previous.rows[0]) == _normalized_header(table.rows[0])
        if consecutive and same_header:
            previous.rows.extend(table.rows[1:])
            previous.page_end = table.page_end or table.page
            previous.continued = True
        else:
            merged.append(table)
    return merged


def _parse_pdf_ocr(file_path: Path) -> ParsedFileContent:
    from app.services.ocr_service import ocr_pdf

    ocr_result = ocr_pdf(file_path)
    pages = [
        ParsedPage(number=page.page, text=_clean_text(page.text))
        for page in ocr_result.pages
    ]
    text = "\n\n".join(page.text for page in pages if page.text)
    warnings: list[str] = []
    if not text:
        warnings.append("OCR 未识别到有效文本")
    elif ocr_result.confidence < 0.75:
        warnings.append(f"OCR 平均置信度较低: {ocr_result.confidence:.2f}")
    return ParsedFileContent(
        text=text,
        pages=pages,
        warnings=warnings,
        is_scanned=True,
        ocr_applied=True,
        ocr_confidence=ocr_result.confidence,
    )


@dataclass(frozen=True)
class DocumentTool:
    name: str
    suffixes: frozenset[str]
    parser: Callable[[Path], ParsedFileContent]
    stage: str = "primary"


class DocumentToolRegistry:
    def __init__(self) -> None:
        self._tools: list[DocumentTool] = []

    def register(self, tool: DocumentTool) -> None:
        self._tools.append(tool)

    def select(self, suffix: str, stage: str = "primary") -> DocumentTool | None:
        return next(
            (
                tool
                for tool in self._tools
                if suffix in tool.suffixes and tool.stage == stage
            ),
            None,
        )

    def capabilities(self) -> list[dict[str, object]]:
        return [
            {
                "name": tool.name,
                "suffixes": sorted(tool.suffixes),
                "stage": tool.stage,
            }
            for tool in self._tools
        ]


document_tool_registry = DocumentToolRegistry()
document_tool_registry.register(
    DocumentTool(
        name="通用文本解析工具",
        suffixes=frozenset({".txt", ".md", ".csv", ".json"}),
        parser=_parse_text,
    )
)
document_tool_registry.register(
    DocumentTool(
        name="XLSX评审数据解析工具",
        suffixes=frozenset({".xlsx"}),
        parser=_parse_xlsx,
    )
)
document_tool_registry.register(
    DocumentTool(
        name="RapidOCR扫描PDF识别工具",
        suffixes=frozenset({".pdf"}),
        parser=_parse_pdf_ocr,
        stage="fallback",
    )
)
document_tool_registry.register(
    DocumentTool(
        name="PDF文本与表格解析工具",
        suffixes=frozenset({".pdf"}),
        parser=_parse_pdf,
    )
)
document_tool_registry.register(
    DocumentTool(
        name="DOCX段落与表格解析工具",
        suffixes=frozenset({".docx"}),
        parser=_parse_docx,
    )
)


def parse_file(path: str | Path) -> ParsedFileContent:
    file_path = Path(path)
    if not file_path.exists():
        raise RuntimeError(f"文件不存在: {file_path}")
    if file_path.stat().st_size == 0:
        raise RuntimeError("文件为空，无法解析")

    suffix = file_path.suffix.lower()
    if suffix == ".doc":
        raise RuntimeError("旧版 .doc 暂不直接支持，请转换为 .docx 后重试")
    tool = document_tool_registry.select(suffix)
    if tool is None:
        raise RuntimeError(f"暂不支持的文件类型: {suffix or '无扩展名'}")

    result = tool.parser(file_path)
    result.selected_tool = tool.name
    result.tool_trace = [f"选择工具: {tool.name}", f"执行工具: {tool.name}"]
    if result.is_scanned:
        fallback_tool = document_tool_registry.select(suffix, stage="fallback")
        if fallback_tool is None:
            result.tool_trace.append("降级策略: 未配置 OCR，标记人工复核")
        else:
            result.tool_trace.append(f"选择后备工具: {fallback_tool.name}")
            try:
                ocr_result = fallback_tool.parser(file_path)
                result.tool_trace.append(f"执行后备工具: {fallback_tool.name}")
                result.ocr_applied = True
                result.ocr_confidence = ocr_result.ocr_confidence
                result.selected_tool = f"{tool.name} -> {fallback_tool.name}"
                if ocr_result.text:
                    result.text = ocr_result.text
                    result.pages = ocr_result.pages
                    result.warnings = ocr_result.warnings
                else:
                    result.warnings.extend(ocr_result.warnings)
            except Exception as exc:
                result.tool_trace.append(f"后备工具失败: {fallback_tool.name}")
                result.warnings.append(f"OCR 执行失败: {exc}")
    return result


def extract_text(path: str | Path) -> str:
    """兼容现有调用方，仅返回解析后的完整文本。"""
    return parse_file(path).text
