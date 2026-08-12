from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.core.config import ensure_data_dirs, settings
from app.schemas.task import TaskRecord, TaskResult


REPORT_AGENT_ORDER = (
    "文档解析智能体",
    "合规审查智能体",
    "数据核验智能体",
    "异常分析智能体",
    "结果复核智能体",
    "报告生成智能体",
)


def issue_is_confirmed(issue) -> bool:
    return issue.final_status == "confirmed_issue"


def issue_needs_review(issue) -> bool:
    return issue.final_status == "human_review"


def public_warning(warning: str) -> str:
    """Hide implementation/vendor errors from formal deliverables."""
    if "Dify" in warning or "Workflow" in warning or "工作流" in warning:
        return "语义增强未完成，已采用确定性解析结果继续核验。"
    return warning


def report_suggestion(issue, report_status: str) -> str:
    suggestion = issue.suggestion or "待补充处置建议"
    if report_status != "正式核验版" or issue_needs_review(issue):
        return suggestion
    if issue.detection_status == "not_detected":
        return "人工复核已完成；请依据复核结论补正缺失的形式要件并留痕。"
    if issue.detection_status == "mismatch":
        return "人工复核已完成；请核实主体信息，修正不一致内容并留痕。"
    if issue.detection_status in {"not_checked", "low_confidence", "uncertain"}:
        return "人工复核已完成；请依据复核结论完成相应补正、整改并留痕。"
    return suggestion.replace("建议人工复核", "建议依据已完成的人工复核结论").replace(
        "请人工复核", "请依据已完成的人工复核结论"
    ).replace("请人工查看", "请依据已完成的人工复核记录核对")


def confidence_rows(issue) -> tuple[tuple[str, str], ...]:
    detection = f"{issue.confidence:.0%}" if issue.detection_status else "不适用"
    evidence = "完整" if issue.evidence or issue.evidence_refs else "不足"
    conclusion = "待人工复核" if issue_needs_review(issue) else "人工已确认/规则已确认"
    return (("检测置信度", detection), ("证据完整度", evidence), ("结论状态", conclusion))


def report_agent_names(result: TaskResult) -> list[str]:
    present = {item.agent for item in result.agent_results}
    return [name for name in REPORT_AGENT_ORDER if name in present]


def _report_data(result: TaskResult) -> dict:
    reports = [
        item.data
        for item in result.agent_results
        if item.agent == "报告生成智能体" and isinstance(item.data, dict)
    ]
    return reports[-1] if reports else {}


def create_markdown_report(task: TaskRecord, result: TaskResult) -> Path:
    ensure_data_dirs()
    report_path = settings.reports_dir / f"{task.task_id}.md"

    lines: list[str] = [
        f"# {task.project_name} 评标智能核验报告",
        "",
        f"- 任务编号: {task.task_id}",
        f"- 项目编号: {task.project_id}",
        f"- 核验类型: {task.check_type}",
        f"- 核验结论: {final_report_conclusion(result)}",
        f"- 报告状态: {_report_data(result).get('report_status', '待复核版')}",
        "",
        "## 自动路由",
        "",
        f"- 路由模式: {result.routing.get('mode', '未记录')}",
        "- 已选智能体: "
        + ", ".join(result.routing.get("selected_agents", [])),
        "- 路由理由: " + "；".join(result.routing.get("reasons", [])),
        "",
        "## 一、核验资料及文档解析结果",
        "",
    ]

    for doc in result.parsed_documents:
        lines.extend(
            [
                f"### {doc.filename}",
                "",
                f"- 文件类型: {doc.file_type}",
                f"- 文本长度: {doc.text_length}",
                f"- 项目名称: {doc.project_name or '未识别'}",
                f"- 招标人: {doc.tenderer or '未识别'}",
                f"- 投标人: {', '.join(doc.bidders) if doc.bidders else '未识别'}",
                f"- 报价: {', '.join(doc.bid_prices) if doc.bid_prices else '未识别'}",
                "",
            ]
        )

    confirmed_issues = [issue for issue in result.issues if issue_is_confirmed(issue)]
    pending_issues = [issue for issue in result.issues if issue_needs_review(issue)]
    lines.extend(
        [
            "## 二、审查事项清单",
            "",
            f"- 明确问题: {len(confirmed_issues)} 项",
            f"- 待人工判断: {len(pending_issues)} 项",
            "",
        ]
    )

    if not result.issues:
        lines.append("未发现需要处理的审查事项。")

    def append_issues(title: str, issues: list) -> None:
        if not issues:
            return
        lines.extend([f"### {title}", ""])
        for index, issue in enumerate(issues, start=1):
            lines.extend(
                [
                    f"#### {title[:-1]} {index}: {issue.issue_type}",
                    "",
                    f"- 问题编号: {issue.issue_id or '未生成'}",
                    f"- 来源智能体: {issue.agent}",
                    f"- 自动判断: {issue.assessment}",
                    f"- 检测置信度: {issue.confidence:.0%}" if issue.detection_status else "- 检测置信度: 不适用",
                    f"- 证据完整度: {'完整' if issue.evidence or issue.evidence_refs else '不足'}",
                    f"- 结论状态: {'待人工复核' if issue_needs_review(issue) else '人工已确认/规则已确认'}",
                    f"- 风险等级: {issue.risk_level}",
                    f"- 来源文件: {issue.source_file or '未定位'}",
                    f"- 位置: {issue.source_location or '未定位'}",
                    f"- 问题描述: {issue.description}",
                    f"- 依据: {issue.basis or '待人工补充'}",
                    f"- 建议: {report_suggestion(issue, _report_data(result).get('report_status', '待复核版'))}",
                    "",
                ]
            )
            if issue.evidence_refs:
                lines.append("- 可定位证据:")
                for ref in issue.evidence_refs:
                    location = ref.section or "未识别章节"
                    if ref.page:
                        location += f"，第{ref.page}页"
                    lines.append(
                        f"  - [{issue.source_file or ref.document_id} / {location}] {ref.quote}"
                    )
                lines.append("")

    append_issues("明确问题", confirmed_issues)
    append_issues("待人工判断项", pending_issues)

    for agent_result in result.agent_results:
        if agent_result.agent != "人工复核节点":
            continue
        normal_clauses = agent_result.data.get("normal_clauses", [])
        if normal_clauses:
            lines.extend(["### 人工确认的正常条款", ""])
            for item in normal_clauses:
                lines.append(
                    f"- {item.get('evidence') or item.get('description') or item.get('issue_id')}"
                )
            lines.append("")

    lines.extend(["## 三、专项智能体结论", ""])
    for agent_name in report_agent_names(result):
        lines.extend([f"### {agent_name}", "", final_agent_conclusion(result, agent_name), ""])

    human_review = next(
        (
            agent_result
            for agent_result in result.agent_results
            if agent_result.agent == "人工复核节点"
        ),
        None,
    )
    lines.extend(
        [
            "## 四、人工复核意见",
            "",
            (
                human_review.summary
                if human_review
                else "本任务未触发高风险人工复核。"
            ),
            "",
        ]
    )

    lines.extend(["## 五、整改建议汇总", ""])
    report_status = str(_report_data(result).get("report_status", "待复核版"))
    suggestions = list(dict.fromkeys(report_suggestion(issue, report_status) for issue in result.issues))
    if suggestions:
        for index, suggestion in enumerate(suggestions, start=1):
            lines.append(f"{index}. {suggestion}")
    else:
        lines.append("当前没有需要输出的整改建议。")
    lines.extend(["", "## 六、综合结论", "", final_report_conclusion(result), ""])

    report_path.write_text("\n".join(lines), encoding="utf-8")
    return report_path


def _set_run_font(run, size: float = 11, bold: bool | None = None) -> None:
    """standard_business_brief; named CJK fallback override: Microsoft YaHei."""
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def _set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    table._enterprise_widths = list(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=False):
            _set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _add_page_field(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend((begin, instruction, end))
    _set_run_font(run, size=8.5)


def _add_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    _set_run_font(paragraph.add_run(str(text)), size=9.5, bold=bold)


def _configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "招投标智能核验 | 审查报告"
    _set_run_font(header.runs[0], size=9)
    header.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(footer.add_run("评标智能核验报告  |  第 "), size=8.5)
    _add_page_field(footer)
    _set_run_font(footer.add_run(" 页  |  结论按复核状态使用"), size=8.5)


def _add_table_header(table, labels: tuple[str, ...], fill: str = "E8EEF5") -> None:
    _repeat_header(table.rows[0])
    for cell, label in zip(table.rows[0].cells, labels, strict=True):
        _set_cell_fill(cell, fill)
        _add_cell_text(cell, label, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_label_value_table(document, rows: list[tuple[str, object]]) -> None:
    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    _set_table_geometry(table, [2700, 6660])
    for row, (label, value) in zip(table.rows, rows, strict=True):
        _set_cell_fill(row.cells[0], "F2F4F7")
        _add_cell_text(row.cells[0], str(label), bold=True)
        _add_cell_text(row.cells[1], str(value))


def _report_package(result: TaskResult) -> dict:
    data = _report_data(result)
    package = data.get("report_package", {})
    return package if isinstance(package, dict) else {}


def select_report_issues(result: TaskResult, report_type: str) -> list:
    """Apply the user's output-document scope instead of always exporting all issues."""
    agent_by_type = {
        "合规审查专项报告": "合规审查智能体",
        "数据核验专项报告": "数据核验智能体",
        "异常分析专项报告": "异常分析智能体",
    }
    selected_agent = agent_by_type.get(report_type)
    issues = list(result.issues)
    if selected_agent:
        issues = [issue for issue in issues if issue.agent == selected_agent]
    risk_order = {"高": 0, "中": 1, "低": 2}
    return sorted(issues, key=lambda item: (risk_order.get(item.risk_level, 9), item.issue_id))


def final_agent_conclusion(result: TaskResult, agent_name: str, report_issues: list | None = None) -> str:
    """Build report prose from the final reviewed issue list, never from stale raw-agent counts."""
    issues = list(result.issues if report_issues is None else report_issues)
    agent_issues = [issue for issue in issues if issue.agent == agent_name]
    confirmed = sum(1 for issue in agent_issues if issue_is_confirmed(issue))
    pending = sum(1 for issue in agent_issues if issue_needs_review(issue))
    documents = result.parsed_documents
    file_count = len(documents)
    page_count = sum(int(getattr(item, "page_count", 0) or 0) for item in documents)

    if agent_name == "文档解析智能体":
        base = f"完成{file_count}份文档"
        if page_count:
            base += f"共{page_count}页"
        base += "的内容解析，并提取正文、章节、表格及关键业务字段。"
        if pending:
            return base + f"视觉或解析质量核验发现{pending}项待人工复核事项，已纳入最终问题清单。"
        return base + "当前未形成需要输出的文档解析质量问题。"
    if agent_name == "合规审查智能体":
        if confirmed or pending:
            return f"完成合规规则核验；经统一结果复核，形成明确问题{confirmed}项、待人工复核事项{pending}项。"
        return "完成合规规则核验；经统一结果复核，当前未形成需要输出的明确或潜在合规问题。"
    if agent_name == "数据核验智能体":
        if confirmed or pending:
            return f"完成关键字段及相关数据交叉核验；经统一结果复核，形成明确问题{confirmed}项、待人工复核事项{pending}项。"
        return "完成关键字段及相关数据交叉核验，当前未发现具有充分证据支持的数据一致性问题。"
    if agent_name == "异常分析智能体":
        if confirmed or pending:
            return f"完成跨文件、跨主体及多信号关联分析；经统一结果复核，形成明确异常{confirmed}项、待人工复核线索{pending}项。"
        scope = "当前材料" if file_count != 1 else "基于当前单份材料"
        return f"{scope}暂未发现可形成有效证据链的异常关联线索。"
    if agent_name == "结果复核智能体":
        return f"已统一复核各专项智能体发现，最终确认明确问题{sum(issue_is_confirmed(i) for i in issues)}项、待人工复核事项{sum(issue_needs_review(i) for i in issues)}项。"
    raw = next((item.summary for item in result.agent_results if item.agent == agent_name), "")
    return raw or "已完成本环节处理。"


def final_report_conclusion(result: TaskResult, report_issues: list | None = None) -> str:
    """One canonical conclusion used by every report section and export format."""
    issues = list(result.issues if report_issues is None else report_issues)
    confirmed = sum(1 for issue in issues if issue_is_confirmed(issue))
    pending = sum(1 for issue in issues if issue_needs_review(issue))
    if not issues:
        return "本次自动核验暂未发现证据充分的明确问题，亦无待人工复核事项。"
    if confirmed == 0:
        return f"本次自动核验暂未发现证据充分的明确问题，发现{pending}项待人工复核事项，建议核对相关原始材料后形成最终结论。"
    if pending == 0:
        return f"本次自动核验形成明确问题{confirmed}项，未发现待人工复核事项，请按问题清单完成整改与闭环确认。"
    return f"本次自动核验形成明确问题{confirmed}项、待人工复核事项{pending}项；请优先处置明确问题，并核对相关原始材料后确认待复核事项。"


def report_display_title(report_type: str, template_type: str) -> str:
    if report_type == "标准化评标报告" or template_type == "标准化评标报告":
        return "标准化评标报告"
    if report_type == "整改建议报告" or template_type == "整改建议报告":
        return "评标核验整改建议报告"
    if template_type == "简版管理层报告":
        return f"{report_type}（管理层简版）"
    if template_type == "详细审查报告":
        return f"{report_type}（详细版）"
    return report_type


def create_docx_report(task: TaskRecord, result: TaskResult) -> Path:
    """Generate an enterprise-grade, traceable evaluation verification report."""
    ensure_data_dirs()
    report_path = settings.reports_dir / f"{task.task_id}.docx"
    document = Document()
    _configure_docx(document)
    report_data = _report_data(result)
    package = _report_package(result)
    project_info = package.get("project_info", {}) if isinstance(package, dict) else {}
    now = datetime.now().astimezone()
    report_status = str(report_data.get("report_status", "待复核版"))
    report_type = str(report_data.get("output_type", "综合智能核验报告"))
    template_type = str(report_data.get("template_type", "标准审查报告"))
    report_title = report_display_title(report_type, template_type)
    report_issues = select_report_issues(result, report_type)
    if template_type == "简版管理层报告":
        report_issues = report_issues[:10]
    counts = {
        level: sum(1 for issue in report_issues if issue.risk_level == level)
        for level in ("高", "中", "低")
    }
    pending_count = sum(1 for issue in report_issues if issue_needs_review(issue))

    # Editorial cover: restrained formal report packaging.
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(72)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    kicker_run = kicker.add_run("招投标全过程智能核验")
    _set_run_font(kicker_run, size=11, bold=True)
    kicker_run.font.color.rgb = RGBColor.from_string("7A5A00")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    title_run = title.add_run(report_title)
    _set_run_font(title_run, size=28, bold=True)
    title_run.font.color.rgb = RGBColor.from_string("0B2545")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(56)
    subtitle_run = subtitle.add_run(project_info.get("project_name") or task.project_name)
    _set_run_font(subtitle_run, size=15, bold=True)
    subtitle_run.font.color.rgb = RGBColor.from_string("1F4D78")

    _add_label_value_table(
        document,
        [
            ("项目编号", task.project_id or "未提供"),
            ("报告编号", f"{task.task_id}-R01"),
            ("报告类型", report_type),
            ("文档模板", template_type),
            ("报告状态", report_status),
            ("生成日期", now.strftime("%Y年%m月%d日")),
            ("核验范围", f"{len(result.parsed_documents)}份资料、{len(result.agent_results)}个执行节点"),
        ],
    )
    document.add_paragraph("")
    confidentiality = document.add_paragraph()
    confidentiality.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = confidentiality.add_run("内部工作文件 · 未经授权不得外传")
    _set_run_font(run, size=9.5, bold=True)
    run.font.color.rgb = RGBColor.from_string("9B1C1C")
    document.add_page_break()

    # These templates intentionally produce different documents, not merely
    # different labels on the same report.
    if template_type == "简版管理层报告":
        document.add_heading("一、管理层结论", level=1)
        document.add_paragraph(final_report_conclusion(result, report_issues))
        _add_label_value_table(
            document,
            [
                ("高风险", counts["高"]),
                ("中风险", counts["中"]),
                ("低风险", counts["低"]),
                ("待人工复核", pending_count),
                ("本简版展示事项", len(report_issues)),
            ],
        )
        document.add_heading("二、重点风险事项", level=1)
        for index, issue in enumerate(report_issues, start=1):
            document.add_heading(f"{index}. {issue.issue_type}（{issue.risk_level}）", level=2)
            document.add_paragraph(issue.description)
            document.add_paragraph(f"建议：{report_suggestion(issue, report_status)}")
        if not report_issues:
            document.add_paragraph("未发现需要向管理层报告的明确或潜在问题。")
        document.add_heading("三、管理决策建议", level=1)
        document.add_paragraph(
            "高风险事项应在评标结论使用前完成专项复核；中风险事项应在报告定稿前完成资料核对；"
            "所有人工复核结论应留痕并关联原始证据。"
        )
        document.add_heading("确认", level=1)
        document.add_paragraph("负责人：________________    日期：________年____月____日")
        for table in document.tables:
            widths = getattr(table, "_enterprise_widths", None)
            if widths:
                _set_table_geometry(table, widths)
        document.save(report_path)
        return report_path

    if template_type == "整改建议报告":
        document.add_heading("一、整改任务概览", level=1)
        document.add_paragraph(final_report_conclusion(result, report_issues))
        document.add_heading("二、整改任务清单", level=1)
        rectification = document.add_table(rows=1, cols=6)
        rectification.style = "Table Grid"
        _set_table_geometry(rectification, [600, 1100, 1650, 3100, 1200, 1710])
        _add_table_header(rectification, ("序号", "编号", "风险", "整改措施", "优先级", "责任/状态"))
        for index, issue in enumerate(report_issues, start=1):
            priority = "立即" if issue.risk_level == "高" else "优先" if issue.risk_level == "中" else "常规"
            values = (
                index,
                issue.issue_id or "未生成",
                issue.risk_level,
                report_suggestion(issue, report_status),
                priority,
                "待复核" if issue_needs_review(issue) else "待整改",
            )
            cells = rectification.add_row().cells
            for cell, value in zip(cells, values, strict=True):
                _add_cell_text(cell, value)
        document.add_heading("三、整改验收要求", level=1)
        document.add_paragraph(
            "整改完成后应记录责任人、完成时间、修改后内容、复核人和复核结论，并将整改证据与问题编号关联归档。"
        )
        document.add_heading("整改确认", level=1)
        document.add_paragraph("整改责任人：________________    复核人：________________")
        document.add_paragraph("完成日期：________年____月____日")
        for table in document.tables:
            widths = getattr(table, "_enterprise_widths", None)
            if widths:
                _set_table_geometry(table, widths)
        document.save(report_path)
        return report_path

    document.add_heading("文档控制信息", level=1)
    _add_label_value_table(
        document,
        [
            ("版本", "V1.0"),
            ("编制方式", "规则引擎、智能体分析与人工复核结果综合生成"),
            ("编制时间", now.strftime("%Y-%m-%d %H:%M:%S %z")),
            ("复核状态", "已完成流程复核" if report_status == "正式核验版" else "存在待复核事项"),
            ("适用范围", "采购文件、响应文件、开评标资料及相关业务记录的辅助核验"),
        ],
    )
    document.add_heading("使用与责任声明", level=2)
    notice = document.add_table(rows=1, cols=1)
    notice.style = "Table Grid"
    _set_table_geometry(notice, [9360])
    _set_cell_fill(notice.cell(0, 0), "FFF8E8")
    _add_cell_text(
        notice.cell(0, 0),
        "本报告用于辅助评审和质量控制，不替代评标委员会、采购人及相关专业人员的法定职责。"
        "标记为待复核、低置信度或证据不足的事项，必须结合原始资料确认后使用。",
    )
    document.add_heading("目录", level=1)
    for item in (
        "一、执行摘要", "二、项目与核验范围", "三、核验方法与执行过程",
        "四、文档解析与数据质量", "五、风险与问题总览", "六、问题逐项说明",
        "七、专项智能体结论", "八、人工复核与整改闭环", "九、综合结论",
        "附录A 证据索引", "附录B 执行与版本信息",
    ):
        document.add_paragraph(item, style="Normal")
    document.add_page_break()

    document.add_heading("一、执行摘要", level=1)
    document.add_paragraph(final_report_conclusion(result, report_issues))
    confirmed_count = sum(1 for issue in report_issues if issue_is_confirmed(issue))
    overview = document.add_table(rows=2, cols=6)
    overview.style = "Table Grid"
    _set_table_geometry(overview, [1560, 1560, 1560, 1560, 1560, 1560])
    for index, label in enumerate(("明确问题", "待人工复核", "高风险", "中风险", "低风险", "合计")):
        _set_cell_fill(overview.rows[0].cells[index], "F2F4F7")
        _add_cell_text(overview.rows[0].cells[index], label, bold=True)
    for index, value in enumerate((confirmed_count, pending_count, counts["高"], counts["中"], counts["低"], len(report_issues))):
        _add_cell_text(overview.rows[1].cells[index], str(value), bold=True)

    significant = [issue for issue in report_issues if issue.risk_level == "高"]
    document.add_heading("重大风险提示", level=2)
    if significant:
        for issue in significant:
            document.add_paragraph(
                f"{issue.issue_type}：{issue.description}", style="List Bullet"
            )
    else:
        document.add_paragraph("本次未形成高风险问题；中低风险事项仍应按整改清单处理。")

    document.add_heading("二、项目与核验范围", level=1)
    _add_label_value_table(
        document,
        [
            ("项目名称", project_info.get("project_name") or task.project_name),
            ("项目编号", task.project_id),
            ("采购人/招标人", project_info.get("tenderer") or "未识别"),
            ("采购代理机构", project_info.get("procurement_agency") or "未识别"),
            ("核验任务类型", task.check_type),
            ("业务系统基准数据", "已提供" if task.system_record else "未提供"),
        ],
    )
    field_sources = project_info.get("field_sources", {})
    if isinstance(field_sources, dict) and field_sources:
        document.add_heading("关键字段来源", level=2)
        source_table = document.add_table(rows=1, cols=6)
        source_table.style = "Table Grid"
        _set_table_geometry(source_table, [1300, 1800, 1900, 1500, 900, 1960])
        _add_table_header(source_table, ("字段", "提取值", "来源文件", "页码/位置", "置信度", "原文"))
        field_labels = {
            "project_name": "项目名称", "tenderer": "采购人/招标人",
            "procurement_agency": "采购代理机构", "budget": "项目预算",
            "price_limit": "最高投标限价", "deadline": "截止时间",
        }
        for name, label in field_labels.items():
            source = field_sources.get(name)
            if not isinstance(source, dict):
                continue
            confidence = float(source.get("confidence", 0))
            value = "待人工确认" if source.get("requires_human_review") or confidence < 0.75 else source.get("value", "")
            values = (
                label, value, source.get("source_file", "未定位"),
                source.get("source_location", "未定位"), f"{confidence:.0%}",
                source.get("source_text", "") or "未记录",
            )
            cells = source_table.add_row().cells
            for cell, cell_value in zip(cells, values, strict=True):
                _add_cell_text(cell, cell_value)
    if template_type == "详细审查报告":
        document.add_paragraph(
            "模板说明：本报告采用详细审查结构，保留执行方法、逐项证据、专项智能体结论、整改闭环及执行版本信息。"
        )
    if template_type == "标准化评标报告":
        document.add_heading("评审结果与中标候选人", level=2)
        rankings = [item for parsed in result.parsed_documents for item in parsed.candidate_rankings]
        if rankings:
            ranking_table = document.add_table(rows=1, cols=4)
            ranking_table.style = "Table Grid"
            _set_table_geometry(ranking_table, [900, 3600, 1300, 3560])
            _add_table_header(ranking_table, ("排名", "候选人", "标段", "原始依据"))
            for item in rankings:
                cells = ranking_table.add_row().cells
                for cell, value in zip(cells, (item.rank, item.bidder, item.lot or "未标注", item.evidence), strict=True):
                    _add_cell_text(cell, value)
        else:
            document.add_paragraph("未从核验资料中识别到可确认的中标候选人排序，需由评标委员会补充确认。")
    document.add_heading("核验资料清单", level=2)
    if not result.parsed_documents:
        document.add_paragraph("未提供可解析的核验资料。")
    else:
        materials = document.add_table(rows=1, cols=7)
        materials.style = "Table Grid"
        _set_table_geometry(materials, [500, 2560, 1300, 900, 800, 800, 2500])
        _add_table_header(
            materials,
            ("序号", "文件名称", "文件类型", "状态", "页数", "表格", "质量提示"),
        )
        for index, parsed in enumerate(result.parsed_documents, start=1):
            cells = materials.add_row().cells
            for cell, value in zip(
                cells,
                (
                    str(index),
                    parsed.filename,
                    parsed.document_subtype or parsed.file_type,
                    parsed.parse_status,
                    parsed.page_count,
                    len(parsed.tables),
                    "；".join(dict.fromkeys(public_warning(item) for item in parsed.warnings[:3])) or "无",
                ),
                strict=True,
            ):
                _add_cell_text(cell, value)

    document.add_heading("三、核验方法与执行过程", level=1)
    document.add_paragraph(
        "本次核验采用确定性规则、文档解析工具、知识库检索、语言模型分析和人工复核相结合的方式。"
        "确定性规则负责字段比对、计算复算和证据约束；智能体负责语义识别、风险解释和报告组织。"
    )
    method_table = document.add_table(rows=1, cols=4)
    method_table.style = "Table Grid"
    _set_table_geometry(method_table, [1900, 2400, 2860, 2200])
    _add_table_header(method_table, ("执行环节", "主要输入", "核验重点", "输出"))
    method_rows = (
        ("文档解析", "采购及评审资料", "OCR、章节、表格、关键字段、印章签名线索", "统一结构化数据"),
        ("合规审查", "条款、评标报告、法规知识库", "完整性、限制性条款、废标依据、法规引用", "合规问题清单"),
        ("数据核验", "报价、评分、排名和系统记录", "合计、权重、平均值、字段及排序一致性", "差异与计算问题"),
        ("异常分析", "多文件、多主体及关系数据", "评分偏离、文件雷同、报价规律和关联信号", "异常线索"),
        ("报告生成", "全部核验结果和人工复核", "去重、复核覆盖、风险统计及整改闭环", "Word/PDF报告"),
    )
    for values in method_rows:
        cells = method_table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            _add_cell_text(cell, value)

    document.add_heading("四、文档解析与数据质量", level=1)
    quality = document.add_table(rows=1, cols=7)
    quality.style = "Table Grid"
    _set_table_geometry(quality, [2100, 900, 900, 900, 900, 900, 2760])
    _add_table_header(quality, ("文件", "章节", "表格", "OCR", "印章检查", "状态", "告警摘要"))
    for parsed in result.parsed_documents:
        seal_state = "已检查" if parsed.seal_signature_checks else "未形成结果"
        values = (
            parsed.filename,
            len(parsed.sections),
            len(parsed.tables),
            "已启用" if parsed.ocr_applied else "未启用",
            seal_state,
            parsed.parse_status,
            "；".join(dict.fromkeys(public_warning(item) for item in parsed.warnings[:5])) or "无",
        )
        cells = quality.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            _add_cell_text(cell, value)

    document.add_heading("五、风险与问题总览", level=1)
    if report_issues:
        matrix = document.add_table(rows=1, cols=7)
        matrix.style = "Table Grid"
        _set_table_geometry(matrix, [700, 1200, 1400, 2900, 950, 1000, 1210])
        _add_table_header(matrix, ("序号", "编号", "问题类型", "问题摘要", "风险", "来源", "复核状态"))
        for index, issue in enumerate(report_issues, start=1):
            values = (
                index,
                issue.issue_id or "未生成",
                issue.issue_type,
                issue.description,
                issue.risk_level,
                issue.agent,
                "待复核" if issue_needs_review(issue) else "明确问题",
            )
            cells = matrix.add_row().cells
            for cell, value in zip(cells, values, strict=True):
                _add_cell_text(cell, value)
    else:
        document.add_paragraph("未发现需要输出的明确或潜在问题。")

    document.add_heading("六、问题逐项说明", level=1)
    if not report_issues:
        document.add_paragraph("未发现需要输出的明确或潜在问题。")
    for index, issue in enumerate(report_issues, start=1):
        document.add_heading(f"{index}. {issue.issue_type}（{issue.risk_level}）", level=2)
        evidence = "\n".join(issue.evidence) or "\n".join(
            ref.quote for ref in issue.evidence_refs
        ) or "未提供"
        rows = [
            ("问题编号", issue.issue_id or "未生成"),
            ("来源智能体", issue.agent),
            ("风险等级", issue.risk_level),
            ("来源文件", issue.source_file or "未定位"),
            ("证据位置", issue.source_location or "未定位"),
            ("问题描述", issue.description),
            ("原文证据", evidence),
            ("判断依据", issue.basis or "待人工补充"),
            ("修改建议", report_suggestion(issue, report_status)),
            ("最终状态", "待人工复核" if issue_needs_review(issue) else "明确问题"),
            ("检测状态", issue.detection_status or "不适用"),
            *confidence_rows(issue),
            ("人工复核", "需要" if issue_needs_review(issue) else "已完成或无需复核"),
        ]
        _add_label_value_table(document, rows)

    document.add_heading("七、专项智能体结论", level=1)
    for agent_name in report_agent_names(result):
        document.add_heading(agent_name, level=2)
        document.add_paragraph(final_agent_conclusion(result, agent_name, report_issues))

    document.add_heading("八、人工复核与整改闭环", level=1)
    human_review = next(
        (
            item
            for item in result.agent_results
            if item.agent == "人工复核节点"
        ),
        None,
    )
    if human_review:
        document.add_paragraph(human_review.summary)
        reviewer = str(human_review.data.get("reviewer", "")).strip()
        if reviewer:
            document.add_paragraph(f"复核人：{reviewer}")
    else:
        document.add_paragraph("本任务尚未形成已提交的人工复核结论。")
    suggestions = list(dict.fromkeys(report_suggestion(issue, report_status) for issue in report_issues))
    if suggestions:
        rectification = document.add_table(rows=1, cols=5)
        rectification.style = "Table Grid"
        _set_table_geometry(rectification, [650, 1450, 3860, 1400, 2000])
        _add_table_header(rectification, ("序号", "问题编号", "整改措施", "优先级", "状态/责任建议"))
        for index, issue in enumerate(report_issues, start=1):
            priority = "立即" if issue.risk_level == "高" else "优先" if issue.risk_level == "中" else "常规"
            values = (
                index,
                issue.issue_id or "未生成",
                report_suggestion(issue, report_status),
                priority,
                "待人工复核" if issue_needs_review(issue) else "纳入整改跟踪",
            )
            cells = rectification.add_row().cells
            for cell, value in zip(cells, values, strict=True):
                _add_cell_text(cell, value)
    else:
        document.add_paragraph("当前没有需要输出的整改建议。")

    document.add_heading("九、综合结论", level=1)
    document.add_paragraph(final_report_conclusion(result, report_issues))
    document.add_paragraph(
        "处置建议：对高风险事项应优先暂停相关结论的直接使用并组织专项复核；"
        "对中风险事项应在报告定稿前完成资料核对和修正；低风险事项应纳入文档规范化整改。"
    )

    document.add_page_break()
    document.add_heading("附录A 证据索引", level=1)
    evidence_table = document.add_table(rows=1, cols=6)
    evidence_table.style = "Table Grid"
    _set_table_geometry(evidence_table, [850, 1100, 1900, 1450, 1200, 2860])
    _add_table_header(evidence_table, ("证据编号", "问题编号", "文件", "页码/位置", "证据类型", "说明"))
    evidence_index = 0
    for issue in report_issues:
        refs = issue.evidence_refs
        if refs:
            for ref in refs:
                evidence_index += 1
                location = ref.section or "未定位"
                if ref.page:
                    location += f"，第{ref.page}页"
                evidence_type = {
                    "text": "原文证据", "table": "表格证据", "metadata": "元数据", "derived": "检测结果"
                }.get(ref.source_type, ref.source_type)
                values = (f"E-{evidence_index:03d}", issue.issue_id or "未生成", issue.source_file or ref.document_id, location, evidence_type, ref.quote)
                cells = evidence_table.add_row().cells
                for cell, value in zip(cells, values, strict=True):
                    _add_cell_text(cell, value)
        else:
            for quote in issue.evidence or ["未提供"]:
                evidence_index += 1
                values = (f"E-{evidence_index:03d}", issue.issue_id or "未生成", issue.source_file or "未定位", issue.source_location or "未定位", "原文证据", quote)
                cells = evidence_table.add_row().cells
                for cell, value in zip(cells, values, strict=True):
                    _add_cell_text(cell, value)

    document.add_heading("附录B 执行与版本信息", level=1)
    execution_rows = []
    for agent_result in result.agent_results:
        execution_rows.append(
            (
                agent_result.agent,
                agent_result.data.get("workflow_version", "本地规则/未记录"),
                agent_result.data.get("ruleset_version", "未记录"),
                agent_result.data.get("execution_mode", "未记录"),
            )
        )
    execution = document.add_table(rows=1, cols=4)
    execution.style = "Table Grid"
    _set_table_geometry(execution, [2300, 2200, 2200, 2660])
    _add_table_header(execution, ("执行节点", "工作流版本", "规则版本", "执行模式"))
    for values in execution_rows:
        cells = execution.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            _add_cell_text(cell, value)

    document.add_heading("确认与签署", level=1)
    document.add_paragraph("评标委员会确认意见：________________________________________________________")
    document.add_paragraph("评标委员会负责人：________________    日期：________年____月____日")
    document.add_paragraph("采购人/招标人确认：________________    日期：________年____月____日")
    for table in document.tables:
        widths = getattr(table, "_enterprise_widths", None)
        if widths:
            _set_table_geometry(table, widths)
    document.save(report_path)
    return report_path


def create_reports(task: TaskRecord, result: TaskResult) -> dict[str, Path]:
    from app.services.pdf_service import create_report_pdf

    return {
        "markdown": create_markdown_report(task, result),
        "docx": create_docx_report(task, result),
        "pdf": create_report_pdf(task, result),
    }

